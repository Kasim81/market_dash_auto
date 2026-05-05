"""
Build `Data_Pipeline_Project_Report.docx`.

This is the final assembled report. Sections 1–13, Appendices A–B,
title page, table of contents, page header, page numbers in the
footer, and document metadata. The interim equation-library probe
appendix, present from Steps 2 through 5 for typography review, has
been removed; every equation lives in its final position in the body.

Build steps (chronological — for repository history reference):
  1. Skeleton + OMML helper module                     [committed]
  2. Equation library (25 equations) + probe page      [committed]
  3. Sections 1–5 (overview / architecture / patterns) [committed]
  4. Sections 6–8 (Phase E / regime / ops scaffolding) [committed]
  5. Sections 9–13 + Appendices A–B                    [committed]
  6. TOC + headers / footers + remove probe            [this commit]

Usage
-----
    cd manuals
    python build_project_report.py
"""

import os
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Sibling-module imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import _omml as M           # noqa: E402
import _equations as EQ     # noqa: E402


OUT_FILENAME = "Data_Pipeline_Project_Report.docx"
REPORT_DATE  = "5 May 2026"


# Colour palette (consistent across the document)
COLOR_HEADER     = "D9E2F3"   # slate-blue header rows
COLOR_CALLOUT    = "F2F2F2"   # light-grey callout fill
COLOR_GREEN      = "DEEBD8"   # status: live / fresh
COLOR_AMBER      = "FFF2CC"   # status: partial / stale
COLOR_PINK       = "FBDDDD"   # status: missing / gap
COLOR_DIAGRAM    = "EAF1FB"   # data-flow box fill


# ── document setup ──────────────────────────────────────────────────────

def setup_document(doc):
    """Apply the report's default page geometry and base typography."""
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)


def add_title_page(doc):
    """Centred title block. Page-break appended."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Market Dashboard Auto")
    r.bold = True
    r.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Project Report — Data Pipeline & Composite Indicator Layer")
    r.font.size = Pt(14)

    for _ in range(3):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Report dated {REPORT_DATE}")
    r.italic = True
    r.font.size = Pt(11)

    doc.add_page_break()


# ── document metadata ───────────────────────────────────────────────────

def set_document_properties(doc):
    """Set Word document core properties (title, subject, author, etc.)."""
    props = doc.core_properties
    props.title    = "Market Dashboard Auto — Project Report"
    props.subject  = ("Data pipeline, composite-indicator layer, "
                      "and operational scaffolding")
    props.author   = "Market Dashboard Auto"
    props.keywords = ("macro market indicators; regime classification; "
                      "z-score; FRED; OECD; Phase E; library-driven")


# ── header / footer / TOC ───────────────────────────────────────────────

def _add_page_field(run, instr_text):
    """Append a Word field expression to a run.

    Used to inject auto-updating fields like PAGE and TOC. The
    runtime sequence is: fldChar(begin) → instrText → fldChar(separate)
    → placeholder text → fldChar(end). Word evaluates the field on
    open and replaces the placeholder.
    """
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    fld_instr = OxmlElement("w:instrText")
    fld_instr.set(qn("xml:space"), "preserve")
    fld_instr.text = instr_text

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(fld_instr)
    run._r.append(fld_end)


def _add_toc_field(paragraph, instr_text, placeholder):
    """Inject a TOC field with a placeholder (Word populates on open)."""
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    fld_instr = OxmlElement("w:instrText")
    fld_instr.set(qn("xml:space"), "preserve")
    fld_instr.text = instr_text

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = placeholder

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(fld_instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def configure_header_footer(doc):
    """Add a page header on every page after the first, plus a centred
    page-number footer. The first page (title page) is left clean.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Header on body pages: small italic title flush right
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("Market Dashboard Auto — Project Report")
    r.italic = True
    r.font.size = Pt(9)

    # Footer on body pages: centred ‘Page X’
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pre = fp.add_run("Page ")
    pre.font.size = Pt(9)
    page_run = fp.add_run()
    page_run.font.size = Pt(9)
    _add_page_field(page_run, "PAGE")


def add_table_of_contents(doc):
    """Insert a TOC field on its own page, after the title page.

    The TOC is generated lazily by Word: on first open Word evaluates
    the field expression and populates the page. The user is prompted
    to update fields if Word's settings require it.
    """
    h = doc.add_heading("Table of Contents", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(8)

    p = doc.add_paragraph()
    _add_toc_field(
        p,
        instr_text='TOC \\o "1-2" \\h \\z \\u',
        placeholder=(
            "Right-click here and choose ‘Update Field’ (or press F9) "
            "to populate the table of contents."
        ),
    )

    doc.add_page_break()


# ── styling helpers ─────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    """Apply a solid fill colour to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def remove_cell_borders(cell):
    """Strip all borders from a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_section_heading(doc, number, title):
    """Top-level numbered section. Starts on a fresh page."""
    doc.add_page_break()
    h = doc.add_heading(f"{number}.  {title}", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(8)


def add_subsection_heading(doc, number, title):
    """Sub-numbered heading."""
    h = doc.add_heading(f"{number}  {title}", level=2)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after  = Pt(2)


def add_paragraph(doc, text, *, italic=False, bold=False, size=10.5,
                  align=None, space_after=6):
    """Body paragraph with consistent formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.italic = italic
    r.bold   = bold
    r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


def add_callout(doc, heading, body, fill=COLOR_CALLOUT):
    """Single-cell shaded box used for asides / architectural notes."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Heading paragraph inside cell
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run(heading)
    r.bold = True
    r.font.size = Pt(10)
    # Body paragraph
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(body)
    r.font.size = Pt(10)
    doc.add_paragraph()  # vertical breathing room


def add_shaded_table(doc, headers, rows, *, header_fill=COLOR_HEADER,
                     row_fills=None, col_widths=None):
    """Standard table: shaded header row, optional per-row fills.

    `rows`     — list of tuples (each tuple == one body row).
    `row_fills`— optional list[str|None] same length as rows.
    `col_widths` — optional list[Inches] same length as headers.
    """
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    # Header
    for ci, h in enumerate(headers):
        c = t.cell(0, ci)
        c.text = ""
        shade_cell(c, header_fill)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)

    # Body
    for ri, row in enumerate(rows, start=1):
        fill = row_fills[ri - 1] if row_fills else None
        for ci, v in enumerate(row):
            c = t.cell(ri, ci)
            c.text = ""
            if fill:
                shade_cell(c, fill)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(str(v))
            r.font.size = Pt(9.5)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for ri in range(1 + len(rows)):
                t.cell(ri, ci).width = w

    doc.add_paragraph()  # spacing after table


def add_dataflow_diagram(doc, steps):
    """Vertical box-and-arrow diagram, rendered as a 1-column borderless
    table. `steps` is a list[str]; ' → ' arrows are inserted between rows.

    The diagram is purely text + cell-shading — no images, no SVG —
    deliberately, so the document renders identically on every Word build.
    """
    nrows = 2 * len(steps) - 1
    t = doc.add_table(rows=nrows, cols=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, step in enumerate(steps):
        # Box row
        cell = t.cell(2 * idx, 0)
        cell.text = ""
        shade_cell(cell, COLOR_DIAGRAM)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(step)
        r.font.size = Pt(9.5)
        r.bold = True

        # Arrow row (between boxes only)
        if idx < len(steps) - 1:
            arrow_cell = t.cell(2 * idx + 1, 0)
            arrow_cell.text = ""
            remove_cell_borders(arrow_cell)
            p = arrow_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            r = p.add_run("▼")
            r.font.size = Pt(10)

    doc.add_paragraph()


def add_eq_inline(paragraph, builder):
    M.add_inline_equation(paragraph, builder())


def add_eq_display(doc, builder, *, eq_number=None):
    """Centred display equation with optional right-aligned (n.n) label."""
    if eq_number is None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(8)
        M.add_display_equation(p, builder())
        return

    # Two-cell borderless table: equation centred, number right-aligned.
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(5.4)
    t.columns[1].width = Inches(0.9)
    eq_cell, num_cell = t.cell(0, 0), t.cell(0, 1)
    remove_cell_borders(eq_cell)
    remove_cell_borders(num_cell)

    eq_p = eq_cell.paragraphs[0]
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    M.add_display_equation(eq_p, builder())

    num_p = num_cell.paragraphs[0]
    num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = num_p.add_run(f"({eq_number})")
    r.font.size = Pt(10)
    doc.add_paragraph()


# ── indicator-card helper (§6 worked examples) ──────────────────────────

def _card_label_cell(t, ri, text):
    """Left cell — label (small caps, bold)."""
    c = t.cell(ri, 0)
    shade_cell(c, COLOR_CALLOUT)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9.5)


def _card_text_cell(t, ri, text, *, italic=False, size=9.5):
    c = t.cell(ri, 1)
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)


def _card_regime_cell(t, ri, regime_rows):
    """Right cell rendering condition → label pairs as separate lines."""
    c = t.cell(ri, 1)
    first = True
    for cond, label in regime_rows:
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(cond)
        r1.font.size = Pt(9.5)
        r1.bold = True
        r2 = p.add_run("   →   ")
        r2.font.size = Pt(9.5)
        r3 = p.add_run(label)
        r3.font.size = Pt(9.5)


def _card_formula_cell(t, ri, eq_builder):
    """Right cell containing a centred display equation."""
    c = t.cell(ri, 1)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    M.add_display_equation(p, eq_builder())


def add_indicator_card(doc, *, ind_id, full_name, eq_builder, data,
                       lookback, regime_rows, interpretation, references):
    """Render one indicator worked-example as a structured card."""
    rows = 7   # title + 6 field rows
    t = doc.add_table(rows=rows, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    t.columns[0].width = Inches(1.1)
    t.columns[1].width = Inches(5.4)

    # Title row (merged across both columns, slate-blue header)
    title = t.cell(0, 0).merge(t.cell(0, 1))
    shade_cell(title, COLOR_HEADER)
    p = title.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f"{ind_id}  —  {full_name}")
    r.bold = True
    r.font.size = Pt(11)

    _card_label_cell(t, 1, "Formula")
    _card_formula_cell(t, 1, eq_builder)

    _card_label_cell(t, 2, "Data")
    _card_text_cell(t, 2, data)

    _card_label_cell(t, 3, "Lookback")
    _card_text_cell(t, 3, lookback)

    _card_label_cell(t, 4, "Regime")
    _card_regime_cell(t, 4, regime_rows)

    _card_label_cell(t, 5, "Interpretation")
    _card_text_cell(t, 5, interpretation)

    _card_label_cell(t, 6, "References")
    _card_text_cell(t, 6, references, italic=True, size=9)

    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════
#                       SECTION 1 — EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════

def add_section_1(doc):
    add_section_heading(doc, 1, "Executive Summary")

    add_paragraph(doc,
        "Market Dashboard Auto is a fully-automated, free-data macro-market "
        "intelligence pipeline that runs nightly on GitHub Actions and "
        "produces, for each trading day: a snapshot of approximately three "
        "hundred and ninety financial instruments across all major asset "
        "classes; weekly historical price series back to 1950; raw "
        "macro-economic data covering twelve economies and ten public-source "
        "providers back to 1947; and a derived layer of ninety-two composite "
        "indicators, each carrying a rolling z-score, a regime classification, "
        "a forward-regime signal, and a z-score trajectory diagnostic. "
        "Outputs land in Google Sheets, in version-controlled CSV files, and "
        "in an interactive HTML explorer. The pipeline has been operating in "
        "production since the consolidation of its multi-source coordinator "
        "in April 2026, and it is the substrate on which downstream "
        "analytical work — including regime-based asset allocation research — "
        "is being built."
    )

    add_paragraph(doc,
        "The design centres on a single architectural rule: every identifier "
        "the pipeline fetches lives in a CSV registry rather than in Python "
        "code. This rule is enforced by an integrated daily audit covering "
        "fetch outcomes, schema drift across three history–library pairs, "
        "value-change staleness against per-frequency tolerances, and "
        "history preservation under retroactive source-side truncation. "
        "Operational scaffolding — a writeback step that retires dead "
        "tickers automatically after a fourteen-day streak, a perpetual "
        "GitHub Issue thread for the daily heartbeat, and an idempotent "
        "Sheets-tab cleanup mechanism — removes the operator from the "
        "daily loop entirely."
    )

    # Key-facts panel
    add_paragraph(doc, "Key facts at a glance:", bold=True, size=10.5,
                  space_after=2)

    facts = [
        ("Instruments tracked",            "~390 across 8 asset classes"),
        ("Macro series ingested",          "~150 across 12 economies"),
        ("Composite indicators computed",  "92 (90 Leading, 2 Coincident)"),
        ("Free public data sources",       "10 (FRED, OECD, World Bank, IMF, "
                                           "DB.nomics, ifo, BoE, ECB, BoJ, e-Stat)"),
        ("Code base",                      "~9,400 lines of Python across "
                                           "7 top-level modules + 12-module sources/ package"),
        ("Earliest historical data",       "1947 (FRED), 1950 (yfinance), "
                                           "1991 (ifo), 1955 (e-Stat)"),
        ("Daily run wall-clock",           "12–15 minutes end-to-end"),
        ("Operational status",             "Production, ~270 successful daily runs"),
        ("Operator burden",                "Zero per-day; weekly review of GitHub Issue thread"),
    ]
    add_shaded_table(doc,
        headers=["Item", "Value"],
        rows=facts,
        col_widths=[Inches(2.2), Inches(4.3)],
    )

    add_paragraph(doc,
        "This report describes the system as it stands today. Section 2 "
        "frames the problem the project addresses and the constraints "
        "imposed by the free-data stance. Section 3 presents the end-to-end "
        "system in a single data-flow view. Sections 4 and 5 cover the data "
        "architecture and the design principles that make the system "
        "durable. Section 6 documents the composite-indicator methodology "
        "in detail, including six worked examples spanning the indicator "
        "families. Section 7 covers regime classification. Section 8 "
        "describes the operational scaffolding. Sections 9 and 10 cover "
        "the user-facing artefacts and the integration surface. Sections "
        "11 to 13 record current coverage, accepted gaps, and the residual "
        "limitations the system must carry. The mathematical notation is "
        "consolidated in Appendix A; references are in Appendix B."
    )


# ════════════════════════════════════════════════════════════════════════
#                  SECTION 2 — PROBLEM STATEMENT & MOTIVATION
# ════════════════════════════════════════════════════════════════════════

def add_section_2(doc):
    add_section_heading(doc, 2, "Problem Statement & Motivation")

    add_subsection_heading(doc, "2.1", "The macro-market analyst's daily problem")

    add_paragraph(doc,
        "Interpretation of current market conditions requires consolidated, "
        "cross-asset, cross-region context. A practitioner must be able to "
        "see, in a single workspace, where equity rotation is pointing, "
        "what the credit complex is signalling, whether the yield curve "
        "agrees, how volatility regimes are evolving, what the major "
        "macro-economic releases are saying, and whether the supporting "
        "evidence is strengthening or weakening. Commercial platforms — "
        "Bloomberg, FactSet, Refinitiv — assemble this picture at a "
        "professional-licence cost typically in the range of twenty to "
        "thirty thousand pounds sterling per seat per annum. For an "
        "individual practitioner or a small research team, that cost is "
        "prohibitive."
    )

    add_paragraph(doc,
        "The free-data alternative is fragmented. Each major statistical "
        "agency publishes data through its own API, with its own schema, "
        "rate-limit profile, error semantics, and cadence. The Federal "
        "Reserve's FRED service is rich and well-engineered but covers "
        "only United States macro statistics in depth. The OECD's SDMX "
        "interface offers cross-country coverage at the cost of a "
        "verbose query convention and tight rate limits. The Bank of "
        "Japan, the Bank of England, the European Central Bank, the "
        "Statistics Bureau of Japan, the ifo Institute and DB.nomics "
        "each speak different dialects again. yfinance, the de facto "
        "free price feed for non-commercial users, has its own pattern "
        "of brittle behaviour around delisted instruments, regional "
        "exchanges, and currency conventions. None of these sources is "
        "individually sufficient; together, they require careful "
        "integration to be reliable."
    )

    add_subsection_heading(doc, "2.2", "Constraints")

    add_paragraph(doc,
        "Three constraints follow from the practitioner setting and have "
        "shaped every subsequent design decision:"
    )

    add_bullet(doc,
        "Free, openly accessible sources only. The pipeline accepts no paid "
        "feed and no scraped, terms-of-service-violating route. Where a "
        "desired indicator is available only behind a paywall — Caixin "
        "China Manufacturing PMI, S&P Global PMI series, ICE corporate "
        "bond indices in their full historical depth — that indicator is "
        "either replaced with a free proxy or recorded as an accepted gap."
    )
    add_bullet(doc,
        "Daily refresh discipline without operator intervention. The "
        "pipeline must execute every day, recover from transient errors "
        "automatically, surface persistent errors through a non-blocking "
        "channel, and continue producing partial output even when an "
        "individual phase fails. The operator must not be required to "
        "log in to keep the pipeline alive."
    )
    add_bullet(doc,
        "Multi-source ingestion under a single architectural discipline. "
        "Without enforced uniformity across sources, the integration cost "
        "of each new provider rises rather than falls. This is the design "
        "pressure that produced the registry-first architecture documented "
        "in Section 4."
    )

    add_subsection_heading(doc, "2.3", "Resolution approach")

    add_paragraph(doc,
        "The free-source constraint is honoured by enumeration: every "
        "fetched identifier is registered in one of the source-of-truth "
        "CSV files in the data/ directory, and that registry is the only "
        "mechanism by which a new series enters the pipeline. Any "
        "reaching for a string literal in Python code that resembles a "
        "fetched identifier is, by adopted rule, prohibited. The rule "
        "originated in response to two avoidable refactors caused by "
        "exactly that drift; it is now enforced both socially (a "
        "pre-commit checklist for collaborators) and technically "
        "(the daily audit's static-checks section flags any column in a "
        "history file that does not trace to a registry row)."
    )

    add_paragraph(doc,
        "Reliability under daily operation is achieved through phase "
        "isolation and per-series error containment. Each runtime phase "
        "is wrapped in a module-level exception handler such that a "
        "failure in any later phase cannot affect outputs already "
        "written by earlier phases. Inside each phase, every individual "
        "ticker, series, or indicator computation is similarly "
        "wrapped — a single bad observation never propagates to the "
        "rest of the output. These two patterns, layered, make partial "
        "success the default behaviour and full failure the exception."
    )

    add_paragraph(doc,
        "Operator removal is achieved through GitHub Actions for the "
        "execution channel, automatic Git commit-and-push of the "
        "outputs for the persistence channel, and a perpetual GitHub "
        "Issue thread for the notification channel. Every daily run "
        "posts a single Markdown comment to the same issue, with the "
        "first line summarising the result as either ALL CLEAN or "
        "as a count of issues. GitHub's native notification mechanism "
        "delivers this to the operator's email; no SMTP infrastructure "
        "is required."
    )

    add_subsection_heading(doc, "2.4", "Scope")

    add_paragraph(doc,
        "Within scope are: (a) ingestion and curation of the data layer, "
        "(b) computation of the composite-indicator layer including "
        "z-score, regime, forward-regime, and trajectory semantics, (c) "
        "the operational scaffolding required for daily reliability, and "
        "(d) the user-facing presentation layer (Google Sheets tabs and "
        "the interactive HTML explorer). Out of scope at this stage are "
        "alpha generation or signal-following execution, individual-stock "
        "research below the index level, and any analysis requiring "
        "intra-day cadence. Forward-looking analytical work — including "
        "regime-based asset allocation, back-testing, and portfolio "
        "construction — is built on top of this substrate but is not "
        "the subject of this report."
    )


# ════════════════════════════════════════════════════════════════════════
#                       SECTION 3 — SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════════════

def add_section_3(doc):
    add_section_heading(doc, 3, "System Overview")

    add_paragraph(doc,
        "The pipeline is a single waterfall executed once per day. A "
        "GitHub Actions cron job at 00:34 UTC invokes the master "
        "orchestrator, which in turn chains in the historical, "
        "macro-economic, composite-indicator, and explorer-rebuild "
        "phases. After the data phases complete, an integrated audit "
        "runs and a writeback step adjusts the instrument library "
        "based on persistent dead-ticker streaks. Every step's output "
        "is committed back to the repository. The end-to-end wall-clock "
        "runtime is between twelve and fifteen minutes."
    )

    add_subsection_heading(doc, "3.1", "Daily execution flow")

    add_paragraph(doc,
        "The diagram below is the full daily flow. Each stage's exception "
        "handler ensures that downstream failures cannot reach back to "
        "corrupt the outputs of earlier stages."
    )

    add_dataflow_diagram(doc, [
        "GitHub Actions cron — 00:34 UTC",
        "fetch_data.py  •  Simple + Comp pipelines",
        "fetch_hist.py  •  Comp weekly history (1950→)",
        "fetch_macro_economic.py  •  Phase ME (10-source unified registry)",
        "compute_macro_market.py  •  Phase E (92 composite indicators)",
        "docs/build_html.py  •  Indicator Explorer rebuild",
        "data_audit.py  •  Daily integrated audit (4 sections)",
        "audit_writeback.py  •  Library writeback (dead-ticker streaks)",
        "git commit + push  •  Sheets push  •  GitHub Issue comment",
    ])

    add_subsection_heading(doc, "3.2", "Runtime phases")

    add_paragraph(doc,
        "The phase nomenclature reflects the chronological order in which "
        "the project was built rather than the runtime order. Phases A, "
        "B, C and D — the per-source US-FRED, surveys, international, "
        "and business-survey coordinators respectively — were "
        "consolidated into a single Phase ME (Macro-Economic) "
        "coordinator on 23 April 2026; their tabs are now retired and "
        "are swept from the Google Sheets workspace on every run. The "
        "current phase inventory is:"
    )

    rows = [
        ("Simple Pipeline",
         "fetch_data.py",
         "market_data, sentiment_data",
         "Production (frozen)"),
        ("Comp Pipeline",
         "fetch_data.py + fetch_hist.py",
         "market_data_comp, market_data_comp_hist",
         "Production"),
        ("Phase ME — Macro-Economic",
         "fetch_macro_economic.py + sources/*.py",
         "macro_economic, macro_economic_hist",
         "Production"),
        ("Phase E — Macro-Market Indicators",
         "compute_macro_market.py",
         "macro_market, macro_market_hist",
         "Production"),
        ("Phase F — Calculated Fields",
         "absorbed into compute_macro_market.py",
         "(see Phase E outputs)",
         "Mostly absorbed"),
        ("Phase G — Sheets Export Audit",
         "library_utils.py SHEETS_* constants",
         "(no separate output)",
         "Done"),
        ("Phase H — Daily Integrated Audit",
         "data_audit.py + audit_writeback.py",
         "data_audit.txt, audit_comment.md",
         "Done"),
    ]
    add_shaded_table(doc,
        headers=["Phase", "Module(s)", "Output tabs / artefacts", "Status"],
        rows=rows,
        col_widths=[Inches(1.7), Inches(2.0), Inches(1.8), Inches(1.0)],
    )

    add_subsection_heading(doc, "3.3", "Outputs")

    add_paragraph(doc,
        "The pipeline writes seven active Google Sheets tabs, each "
        "mirrored as a CSV in the data/ directory of the repository. "
        "A retired-tab frozenset is swept on every run so that "
        "decommissioned tabs cannot accumulate. In parallel, the "
        "indicator-explorer rebuild produces a self-contained HTML "
        "page and an embedded JavaScript bundle of market data. "
        "All of these artefacts are committed back to the main "
        "branch with a timestamped message; the explorer is "
        "additionally published to GitHub Pages."
    )

    rows = [
        ("market_data",            "Snapshot, ~70 instruments",
         "Simple pipeline; consumed by downstream trigger.py"),
        ("market_data_comp",       "Snapshot, ~390 instruments",
         "Comp pipeline; full instrument library"),
        ("market_data_comp_hist",  "Weekly history from 1950",
         "Friday-spine alignment; 10 metadata prefix rows"),
        ("macro_economic",         "Long-form snapshot",
         "Unified raw-macro layer across all 10 sources"),
        ("macro_economic_hist",    "Wide-form weekly history from 1947",
         "14 metadata prefix rows + Friday spine"),
        ("macro_market",           "92 indicators, snapshot",
         "Includes z-score, regime, fwd_regime, trajectory diagnostics"),
        ("macro_market_hist",      "Weekly indicator history from 2000",
         "raw, zscore, regime, fwd_regime per indicator"),
    ]
    add_shaded_table(doc,
        headers=["Tab / CSV", "Shape", "Contents"],
        rows=rows,
        col_widths=[Inches(1.8), Inches(2.1), Inches(2.6)],
    )

    add_callout(doc,
        "Why a Friday spine?",
        "Macro releases arrive on heterogeneous cadences (daily, weekly, "
        "monthly, quarterly, annual). Aligning every series to a single "
        "weekly Friday-close index allows cross-series analysis and "
        "z-score computation without per-pair date-alignment logic. "
        "Forward-fill into the Friday slots is acceptable because the "
        "downstream value-change staleness audit (Section 8) detects "
        "any series that has stopped moving."
    )


# ════════════════════════════════════════════════════════════════════════
#               SECTION 4 — DATA ARCHITECTURE & SOURCE CATALOGUE
# ════════════════════════════════════════════════════════════════════════

def add_section_4(doc):
    add_section_heading(doc, 4, "Data Architecture & Source Catalogue")

    add_subsection_heading(doc, "4.1", "The single architectural rule")

    add_paragraph(doc,
        "The project's dominant design rule is this: every identifier the "
        "pipeline fetches must live in a CSV registry under data/, never "
        "in Python source. The motivation for the rule is empirical. "
        "Two earlier refactors had to chase hardcoded identifiers out of "
        "the codebase — the indicator-metadata Python dictionary into a "
        "registry CSV in April 2026, and the per-source coordinators into "
        "a unified coordinator one week later. Each refactor cost "
        "engineering time that would not have been incurred had the "
        "identifiers lived in their proper home from the outset."
    )

    add_paragraph(doc,
        "The rule applies to: FRED series identifiers, DB.nomics series "
        "paths, OECD SDMX dataflow and dimension keys, World Bank and IMF "
        "indicator codes, ifo Excel sheet and column locations, ECB and "
        "BoE series codes, BoJ time-series codes, e-Stat statsDataIds, "
        "yfinance tickers, and any URL fragment that selects a specific "
        "dataset. It does not apply to: API base URLs, computation "
        "parameters such as the z-score window, internal column-name "
        "conventions, or calculator wiring (which series each calculator "
        "consumes), since those are logic rather than registry."
    )

    add_callout(doc,
        "Architectural invariant",
        "If a change adds, removes, renames, or substitutes a fetched "
        "identifier, the only file edited is the appropriate "
        "data/macro_library_*.csv (or index_library.csv / "
        "macro_indicator_library.csv depending on the layer). No "
        "exceptions. This invariant is verified daily by the audit's "
        "registry-drift static check."
    )

    add_subsection_heading(doc, "4.2", "The data-layer registry")

    add_paragraph(doc,
        "Fourteen CSV files together constitute the data-layer registry. "
        "They are read at process startup, never written by the runtime "
        "pipeline, and version-controlled in the same repository as the "
        "code. Adding or removing a series is a one-CSV-row commit; "
        "renaming a series is a one-CSV-cell commit."
    )

    rows = [
        ("index_library.csv",            "~401",
         "Comp pipeline tickers (yfinance + FRED bond indices)"),
        ("macro_library_countries.csv",  "12",
         "Country codes + per-source code mappings (WB, IMF, OECD)"),
        ("macro_library_fred.csv",       "~82",
         "FRED series IDs (US + international)"),
        ("macro_library_oecd.csv",       "3",
         "OECD SDMX dataflow + dimension keys"),
        ("macro_library_worldbank.csv",  "1",
         "World Bank WDI indicator codes (CPI YoY)"),
        ("macro_library_imf.csv",        "1",
         "IMF DataMapper indicator codes (real GDP growth)"),
        ("macro_library_dbnomics.csv",   "13",
         "DB.nomics series paths (Eurostat, ISM, IMF/IFS fallbacks)"),
        ("macro_library_ifo.csv",        "26",
         "ifo workbook sheet/column locations"),
        ("macro_library_boe.csv",        "7",
         "BoE IADB series (Bank Rate, SONIA, gilt par/zero curves)"),
        ("macro_library_ecb.csv",        "3",
         "ECB Data Portal SDMX keys (deposit rate, AAA yield curve)"),
        ("macro_library_boj.csv",        "2",
         "BoJ time-series codes (policy rate, Tankan large mfg DI)"),
        ("macro_library_estat.csv",      "1",
         "e-Stat statsDataIds (METI Indices of Industrial Production)"),
        ("source_fallbacks.csv",         "9",
         "Per-indicator T0/T1/T2/T3 fallback chain (documentation-only)"),
        ("macro_indicator_library.csv",  "92",
         "Phase E composite indicator definitions"),
    ]
    add_shaded_table(doc,
        headers=["Registry file", "Rows", "Contents"],
        rows=rows,
        col_widths=[Inches(2.4), Inches(0.6), Inches(3.5)],
    )

    add_subsection_heading(doc, "4.3", "Source catalogue")

    add_paragraph(doc,
        "Ten free public sources currently contribute to the data layer. "
        "Each has its own access protocol, its own rate-limit envelope, "
        "and its own pattern of edge-case behaviour; the per-source "
        "modules in the sources/ package encapsulate these differences "
        "so that the upper layers can treat the data uniformly."
    )

    rows = [
        ("yfinance",     "Yahoo Finance",
         "~390 instruments — equities, ETFs, FX, commodities, crypto, vol",
         "No key"),
        ("FRED",         "St. Louis Fed",
         "~85 series across yields, inflation, labour, credit, surveys",
         "Free key"),
        ("OECD",         "OECD SDMX",
         "Composite Leading Indicator, unemployment, 3-month interbank",
         "No key"),
        ("World Bank",   "World Bank WDI",
         "CPI YoY across 11 economies",
         "No key"),
        ("IMF",          "IMF DataMapper",
         "Real GDP Growth (annual WEO including projections)",
         "No key"),
        ("DB.nomics",    "DB.nomics REST",
         "Eurostat surveys, ISM PMIs, IMF/IFS fallback rates",
         "No key"),
        ("ifo",          "ifo Institute",
         "26 German business-survey series (1991→) via Excel workbook",
         "No key"),
        ("BoE",          "Bank of England IADB",
         "Bank Rate, SONIA, gilt par + zero-coupon yields",
         "No key"),
        ("ECB",          "ECB Data Portal",
         "Deposit Facility Rate, AAA yield curve points",
         "No key"),
        ("BoJ",          "Bank of Japan Time-Series",
         "Policy Rate, Tankan Large Manufacturer DI",
         "No key"),
        ("e-Stat",       "Statistics Bureau Japan",
         "METI Indices of Industrial Production (1955→)",
         "Free App ID"),
    ]
    add_shaded_table(doc,
        headers=["Source", "Provider", "Coverage", "Auth"],
        rows=rows,
        col_widths=[Inches(0.9), Inches(1.5), Inches(3.3), Inches(0.7)],
    )

    add_subsection_heading(doc, "4.4", "Country and currency coverage")

    add_paragraph(doc,
        "The country registry covers twelve economies: the United States, "
        "United Kingdom, Germany, France, Italy, Japan, China, Australia, "
        "Canada, Switzerland, the Eurozone aggregate (EA19/EA20 depending "
        "on the source's convention), and India (added April 2026 to "
        "support the India 10-year bond yield series). The registry is "
        "the single source of truth: the World Bank uses code EMU for "
        "the Eurozone and the IMF DataMapper uses code EURO; both "
        "mappings live in macro_library_countries.csv rather than in "
        "Python."
    )

    add_paragraph(doc,
        "Foreign-exchange coverage spans eighteen currencies, of which "
        "fifteen are indirect-quote (i.e. expressed as foreign currency "
        "per US dollar — JPY, CNY, CAD, INR, KRW, HKD, BRL, TWD, MXN, "
        "ZAR, TRY, IDR, RUB, SAR) and three are direct-quote (GBP, EUR, "
        "AUD). For each non-USD instrument the comp pipeline computes a "
        "USD-adjusted return alongside the local-currency return. The "
        "identity used is:"
    )

    add_eq_display(doc, EQ.usd_return, eq_number="4.1")

    add_paragraph(doc,
        "Indirect-quote currencies require inversion of the FX rate "
        "before applying the identity; the inversion is automatic, "
        "driven by membership in the COMP_FCY_PER_USD set in "
        "library_utils.py."
    )

    add_paragraph(doc,
        "A second pre-processing rule applies to UK-listed instruments. "
        "London Stock Exchange tickers (those ending in .L) report "
        "prices in pence (GBp) rather than pounds (GBP) for some "
        "instruments and not others. Rather than maintaining a "
        "hardcoded list of pence-denominated tickers, the pipeline "
        "applies the dynamic correction:"
    )

    add_eq_display(doc, EQ.pence_correction, eq_number="4.2")

    add_paragraph(doc,
        "The threshold of fifty is empirically chosen: no genuine "
        "GBP-denominated UK security has had a median price below "
        "fifty pounds across the available history, while every "
        "pence-denominated security has."
    )

    add_subsection_heading(doc, "4.5", "Source-fallback chain")

    add_paragraph(doc,
        "Because nine FRED series carrying internationally-mirrored "
        "macro data have stopped updating (the OECD-mirror series for "
        "the Japanese policy rate, the Chinese policy rate, the UK "
        "Bank Rate, Eurozone HICP, German industrial production and "
        "others), the pipeline carries a per-indicator fallback "
        "registry in source_fallbacks.csv. Each row records the canonical "
        "tier-zero source for an indicator together with up to three "
        "fallback sources (T1, T2, T3). The runtime currently realises "
        "the chain implicitly through the read order in the unified "
        "coordinator — later sources overwrite earlier sources at the "
        "column level. Of the nine forcing-function rows, seven have "
        "been resolved through tier-one and tier-two fallbacks (DB.nomics "
        "IMF/IFS, Eurostat, the BoE IADB, the ECB Data Portal, the BoJ "
        "Time-Series API, and e-Stat); two remain accepted gaps because "
        "the underlying Chinese statistical agencies publish no free "
        "programmatic interface."
    )


# ════════════════════════════════════════════════════════════════════════
#                  SECTION 5 — PIPELINE DESIGN PRINCIPLES
# ════════════════════════════════════════════════════════════════════════

def add_section_5(doc):
    add_section_heading(doc, 5, "Pipeline Design Principles")

    add_paragraph(doc,
        "Nine recurring design patterns are used across the codebase. "
        "They are documented here once because each pattern appears in "
        "multiple modules; pattern-by-pattern documentation in each "
        "module would duplicate without adding clarity."
    )

    # 5.1 Phase isolation
    add_subsection_heading(doc, "5.1", "Phase isolation")

    add_paragraph(doc,
        "Every runtime phase after the master orchestrator's main "
        "function is wrapped in a module-level try/except block that "
        "catches every exception, logs it, and proceeds to the next "
        "phase. This pattern is the load-bearing reliability mechanism "
        "of the pipeline. If the macro-economic phase crashes — say, "
        "because OECD SDMX is briefly unavailable — the simple "
        "pipeline, the comp pipeline, and the comp historical phase "
        "have already completed and their outputs are already on disk. "
        "The audit subsequently surfaces the macro-economic failure, "
        "but no data the operator depends on has been lost."
    )

    # 5.2 Per-series try/except
    add_subsection_heading(doc, "5.2", "Per-series exception containment")

    add_paragraph(doc,
        "Inside each phase, every individual fetch — one ticker, one "
        "FRED series, one OECD country fan-out, one indicator "
        "computation — is itself wrapped in a try/except block that "
        "logs the failure and continues. A single bad series never "
        "kills the rest of the output. This is the difference between "
        "‘the pipeline failed’ and ‘the pipeline succeeded with twenty "
        "rows missing’; the latter is a strictly more useful state "
        "because the missing rows are visible in the audit and can be "
        "remediated, while preserved rows continue to inform analysis."
    )

    # 5.3 Friday spine
    add_subsection_heading(doc, "5.3", "Friday-spine alignment")

    add_paragraph(doc,
        "All historical data is aligned to a weekly index of Friday "
        "dates. For a series sampled at any cadence, the value at "
        "Friday t is taken as the most recent observation on or before "
        "t, with forward-fill into Friday slots where the source has "
        "not yet released. Formally:"
    )

    add_eq_display(doc, EQ.friday_spine, eq_number="5.1")

    add_paragraph(doc,
        "The Friday convention is chosen because end-of-week is the "
        "least ambiguous market data point for cross-asset analysis. "
        "The forward-fill is acceptable in conjunction with the "
        "value-change staleness audit (Section 8): any series that "
        "has stopped moving is detected and flagged regardless of "
        "whether its forward-filled cells continue to populate."
    )

    # 5.4 Metadata prefix rows
    add_subsection_heading(doc, "5.4", "Metadata prefix rows")

    add_paragraph(doc,
        "History tabs in Google Sheets carry metadata rows above the "
        "data block. The macro-economic history carries fourteen "
        "prefix rows (Column ID, Series ID, Source, Indicator, Country, "
        "Country Name, Region, Category, Subcategory, Concept, "
        "cycle_timing, Units, Frequency, Last Updated). The market-data "
        "comp history carries ten prefix rows. These rows are visible "
        "to a human reader of the spreadsheet, support the explorer's "
        "by-concept and by-region grouping, and are consistently "
        "skipped by every code path that reads the history file via "
        "an explicit header offset."
    )

    # 5.5 Library-driven configuration
    add_subsection_heading(doc, "5.5", "Library-driven configuration")

    add_paragraph(doc,
        "Beyond the registry rule of Section 4.1, several "
        "configuration knobs that would conventionally live in code "
        "have been moved to CSV columns. The simple-pipeline "
        "instrument selection is a boolean column on the master "
        "library. The validation status of each ticker — CONFIRMED, "
        "PENDING, UNAVAILABLE — is a column. Per-row overrides on "
        "freshness tolerances are columns on the per-source library "
        "files. The result is that operator interventions — disabling "
        "a broken series, widening a tolerance for a slow publisher, "
        "swapping in a proxy ticker — are CSV edits rather than "
        "code commits."
    )

    # 5.6 Diff-check CSV commits
    add_subsection_heading(doc, "5.6", "Diff-check CSV commits")

    add_paragraph(doc,
        "Output CSV files are only committed back to git when their "
        "content has actually changed. This avoids a noisy commit "
        "history during weekends and holidays when markets are closed "
        "and series have not moved, while still committing every "
        "weekday's genuine update. The diff check is byte-level after "
        "a deterministic sort and serialisation, so it is robust to "
        "ordering instabilities in upstream APIs."
    )

    # 5.7 Sheets write pattern
    add_subsection_heading(doc, "5.7", "Sheets write pattern")

    add_paragraph(doc,
        "All writers to Google Sheets follow a five-step pattern: "
        "check the target tab against a protected-tab frozenset and "
        "abort if matched, ensure the tab exists (creating it if "
        "necessary), clear the existing range, write the new data, "
        "and finally sweep a separate legacy-tab frozenset on the "
        "way out so retired tabs cannot accumulate. The two frozensets "
        "— SHEETS_PROTECTED_TABS and SHEETS_LEGACY_TABS_TO_DELETE — "
        "are defined once in library_utils.py and imported by every "
        "writer. Adding or retiring a tab is a one-line edit there; "
        "no other writer needs to know."
    )

    # 5.8 Rate limiting
    add_subsection_heading(doc, "5.8", "Rate limiting and exponential backoff")

    add_paragraph(doc,
        "Every external HTTP call is rate-limited with a per-source "
        "base delay and retried with exponential backoff on transient "
        "failures (HTTP 429 or 5xx). The delays and retry budgets are "
        "tuned per source to its published or empirically-observed "
        "rate envelope:"
    )

    rows = [
        ("yfinance",   "0.3 s",   "—",                            "3"),
        ("FRED",       "0.6 s",   "2, 4, 8, 16, 32 s",            "5"),
        ("OECD",       "4.0 s",   "2, 4, 8, 16, 32 s",            "5"),
        ("World Bank", "1.0 s",   "2, 4, 8, 16, 32 s",            "5"),
        ("IMF",        "1.0 s",   "2, 4, 8, 16, 32 s",            "5"),
        ("ECB",        "0.6 s",   "2, 4, 8, 16 s",                "3"),
        ("BoE",        "0.6 s",   "2, 4, 8, 16 s",                "3"),
        ("BoJ",        "0.6 s",   "2, 4, 8, 16 s",                "3"),
        ("e-Stat",     "0.6 s",   "2, 4, 8, 16 s",                "3"),
    ]
    add_shaded_table(doc,
        headers=["Source", "Base delay", "Backoff schedule", "Max retries"],
        rows=rows,
        col_widths=[Inches(1.3), Inches(1.1), Inches(2.0), Inches(1.2)],
    )

    # 5.9 History preservation
    add_subsection_heading(doc, "5.9", "History preservation under source truncation")

    add_paragraph(doc,
        "Source-side history can shrink retroactively. The cleanest "
        "example is the ICE Data licence change of April 2026, which "
        "obliged FRED to truncate its redistributed ICE BofA spread "
        "series — including BAMLH0A0HYM2, BAMLC0A0CM and "
        "BAMLHE00EHYIOAS — to a rolling three-year window. Twenty "
        "or more years of pre-2023 spread data disappeared from the "
        "FRED side. Without intervention, the next nightly fetch "
        "would have overwritten the local history with the truncated "
        "window, losing the historical depth on which the rolling "
        "z-score depends."
    )

    add_paragraph(doc,
        "The mitigation is architectural. For every primary history "
        "file the pipeline writes, an append-only sister file "
        "captures rows that pre-date the current source-side window. "
        "The detection logic operates per column rather than per "
        "row: a rolling-window source can keep row count constant "
        "while the earliest non-null date walks forward each cycle, "
        "so the test is whether a column's earliest non-null date "
        "in the new fetch is later than the earliest non-null date "
        "previously stored. If so, the rows about to disappear are "
        "appended to the sister file, deduplicated by date, before "
        "the live file is rewritten with the new window. Read paths "
        "transparently union the live and sister files, with the "
        "live file winning where it has a non-null observation."
    )

    add_callout(doc,
        "Out of scope",
        "This rule preserves history that the pipeline already has. "
        "It does not back-fill history that was never captured. Filling "
        "a pre-installation gap would require either a paid ICE Data "
        "subscription or an alternative archive — neither is currently "
        "in scope."
    )


# ════════════════════════════════════════════════════════════════════════
#                SECTION 6 — THE COMPOSITE-INDICATOR LAYER
# ════════════════════════════════════════════════════════════════════════

def add_section_6(doc):
    add_section_heading(doc, 6, "The Composite-Indicator Layer (Phase E)")

    add_paragraph(doc,
        "Phase E is the analytical heart of the project. It consumes the "
        "unified macro-economic history and the comp-pipeline price "
        "history, applies a small library of standard transforms, and "
        "produces ninety-two composite indicators. Each indicator emits, "
        "for every Friday in its history, a raw value, a rolling z-score, "
        "a regime classification, a forward-regime signal, and a z-score "
        "trajectory diagnostic. The indicator is the unit on which all "
        "downstream interpretation is built."
    )

    add_subsection_heading(doc, "6.1", "Computation flow per indicator")

    add_paragraph(doc,
        "Every indicator passes through the same five-stage pipeline. "
        "The first three stages are purely mechanical and identical "
        "across indicators; the last two stages are where economic "
        "judgement enters."
    )

    rows = [
        ("1", "Calculator function",
         "A registered Python function (e.g. _calc_US_G1) reads its "
         "inputs from the unified macro hist or the comp-pipeline price "
         "history, applies the indicator-specific transform, and "
         "returns a weekly Series."),
        ("2", "Rolling z-score",
         "A 156-week (three-year) rolling z-score is computed against "
         "the indicator's own history, with a 52-week minimum warm-up "
         "before the first z-value is emitted."),
        ("3", "Regime classification",
         "A registered classifier function (a lambda keyed by indicator "
         "id in REGIME_RULES) maps the (raw, z) pair to a discrete "
         "regime label."),
        ("4", "Forward regime",
         "The slope of the z-score over the trailing eight weeks is "
         "classified as improving / stable / deteriorating, with a "
         "[leading] suffix where the indicator is naturally leading."),
        ("5", "Z-score trajectory",
         "Sampled at 1-, 4-, and 13-week lags, the z-score's recent "
         "trajectory is classified as intensifying, fading, reversing, "
         "or stable. This is purely diagnostic — it does not feed the "
         "regime call."),
    ]
    add_shaded_table(doc,
        headers=["#", "Stage", "Description"],
        rows=rows,
        col_widths=[Inches(0.4), Inches(1.6), Inches(4.5)],
    )

    add_subsection_heading(doc, "6.2", "The rolling z-score")

    add_paragraph(doc,
        "The rolling z-score is the single normalisation that makes "
        "ninety-two indicators of heterogeneous units comparable. For "
        "an indicator series, the z-score at time t is:"
    )

    add_eq_display(doc, EQ.rolling_zscore, eq_number="6.1")

    add_paragraph(doc,
        "where the window mean and standard deviation are computed "
        "over the trailing W observations:"
    )

    add_eq_display(doc, EQ.zscore_window, eq_number="6.2")

    add_paragraph(doc,
        "The choice of W = 156 weeks corresponds to three calendar "
        "years. Three years is long enough to span the typical short "
        "business cycle while remaining short enough that the window "
        "tracks regime shifts rather than averaging across them. The "
        "minimum-periods threshold of 52 weeks ensures that no "
        "indicator emits a z-score until it has accumulated at least "
        "one year of context, so that early-history values are not "
        "computed against a degenerate sample."
    )

    add_callout(doc,
        "Why z-scores rather than absolute thresholds?",
        "The economic meaning of a credit spread of 500 basis points "
        "in 1995 is materially different from its meaning in 2025. "
        "Spreads have structurally tightened. The same is true of "
        "implied volatility, the yield curve, and most spread relations "
        "in the indicator library. A rolling z-score adapts to the "
        "indicator's own recent regime and answers the only question "
        "that matters in practice: is this reading unusual relative "
        "to the recent history? Where structural levels do matter — "
        "the credit spread above 800 basis points, the inverted yield "
        "curve, the inverted VIX term structure — the regime classifier "
        "applies a level override, as discussed in Section 7."
    )

    add_subsection_heading(doc, "6.3", "Indicator formula templates")

    add_paragraph(doc,
        "Every calculator in the library is a thin wrapper around one of "
        "four standard transforms, plus a small set of derived "
        "operations. The transforms are:"
    )

    add_paragraph(doc, "Log-ratio. Used wherever two prices or two "
                  "yield-class series are compared:", bold=True,
                  size=10.5, space_after=2)
    add_eq_display(doc, EQ.log_ratio, eq_number="6.3")

    add_paragraph(doc, "Composite log-ratio. Used to average several "
                  "log-ratios into a single composite signal:",
                  bold=True, size=10.5, space_after=2)
    add_eq_display(doc, EQ.sum_log_ratio, eq_number="6.4")

    add_paragraph(doc, "Arithmetic difference. Used for yield-curve "
                  "and spread differences:",
                  bold=True, size=10.5, space_after=2)
    add_eq_display(doc, EQ.arith_diff, eq_number="6.5")

    add_paragraph(doc, "Year-on-year change. Used for monthly "
                  "macro-economic series whose level is not directly "
                  "comparable across regimes:",
                  bold=True, size=10.5, space_after=2)
    add_eq_display(doc, EQ.yoy_monthly, eq_number="6.6")

    add_subsection_heading(doc, "6.4", "Indicator families")

    add_paragraph(doc,
        "The ninety-two indicators are organised across regions and "
        "concept families:"
    )

    rows = [
        ("US",                    "34", "Equity rotation, credit, rates, "
                                        "volatility, momentum, macro / survey"),
        ("Europe / UK",           "11", "European cyclicals, EU IG/HY credit, "
                                        "BTP-Bund, gilts, breakeven, GBP credit"),
        ("Japan",                 "1",  "Japan vs global equities"),
        ("Asia (China / India)",  "6",  "Size, growth, rates"),
        ("Global",                "8",  "Risk appetite, EM vs DM, CLI "
                                        "differentials and breadth"),
        ("FX & Commodities",      "9",  "Copper/gold, dollar, iron ore, "
                                        "commodity and FX momentum"),
        ("Survey composites",     "23", "Country-level OECD CLI / business / "
                                        "consumer confidence, ISM, PMI"),
    ]
    add_shaded_table(doc,
        headers=["Family", "Count", "Coverage"],
        rows=rows,
        col_widths=[Inches(2.0), Inches(0.6), Inches(3.9)],
    )

    add_subsection_heading(doc, "6.5", "Six worked examples")

    add_paragraph(doc,
        "The six examples that follow span the full breadth of the "
        "library: an equity-rotation log-ratio, a credit-spread "
        "level-and-z hybrid, a yield-curve arithmetic difference, an "
        "equity-volatility term-structure spread, a cross-commodity "
        "log-ratio, and a multi-country survey composite. Each card "
        "lists the formula, the data sources, the regime thresholds, "
        "an interpretation, and the supporting academic references."
    )

    # ── 1. US_G1 ────────────────────────────────────────────────────────
    add_indicator_card(doc,
        ind_id="US_G1",
        full_name="Cyclicals vs Defensives (Discretionary / Staples)",
        eq_builder=EQ.eq_us_g1,
        data="SPDR Consumer Discretionary ETF (XLY) and SPDR Consumer "
             "Staples ETF (XLP), yfinance total-return.",
        lookback="156-week rolling z-score of the log-ratio.",
        regime_rows=[
            ("z > +1",        "pro-growth — OW discretionary, HY credit, cyclicals"),
            ("|z| ≤ 1",       "neutral — balanced allocation"),
            ("z < −1",        "defensive — OW staples, IG, defensive sectors"),
        ],
        interpretation="Households purchase discretionaries when they "
            "feel financially secure and curtail them in stress; staples "
            "spending is near-inelastic. The ratio therefore distils "
            "aggregate consumer confidence and tends to peak three to "
            "six months ahead of the economic cycle.",
        references="Friedman (1957) permanent-income hypothesis; "
            "Fama & French (1989) on cyclicals leading staples around "
            "cycle turns; Conference Board LEI methodology.",
    )

    # ── 2. US_Cr2 ───────────────────────────────────────────────────────
    add_indicator_card(doc,
        ind_id="US_Cr2",
        full_name="US High-Yield Credit Spread (5-regime)",
        eq_builder=EQ.eq_us_cr2_regime,
        data="ICE BofA US High Yield OAS (FRED: BAMLH0A0HYM2) and the "
             "10-year Treasury yield (FRED: DGS10) for context.",
        lookback="Regime is a level / z-score hybrid; z computed on a "
                 "156-week window.",
        regime_rows=[
            ("Spread > 800 bps   ∨   z > +2",            "opportunity"),
            ("Spread > 500 bps   ∧   z > +1",            "stress"),
            ("400 ≤ spread ≤ 600   ∧   |z| < 1",        "normal"),
            ("Spread < 400 bps   ∧   z < −0.5",         "complacent"),
            ("Spread < 300 bps   ∧   z < −1",           "frothy"),
        ],
        interpretation="The HY spread is the most sensitive real-time "
            "measure of credit conditions. Five regimes are required "
            "because at the extremes, structural levels (the 800 bps "
            "contrarian buy zone, the 300 bps frothy zone) carry signal "
            "that pure z-score normalisation would lose by adapting to "
            "regime drift.",
        references="Altman (1968) Z-score / default-probability "
            "framework; Merton (1974) structural credit model; "
            "Duffie & Singleton (1999); T. Rowe Price post-crisis "
            "recovery study (CFA Institute analysis).",
    )

    # ── 3. US_R1 ────────────────────────────────────────────────────────
    add_indicator_card(doc,
        ind_id="US_R1",
        full_name="Yield-Curve Slope 10Y − 3M",
        eq_builder=EQ.eq_us_r1,
        data="Federal Reserve H.15 release via FRED (T10Y3M = DGS10 − DGS3MO), "
             "in percentage points.",
        lookback="Level-based regime with z-score modulation; z computed "
                 "on 156 weeks.",
        regime_rows=[
            ("Spread < 0",                       "recession-watch"),
            ("Spread > 0   ∧   z > +1",          "early-cycle"),
            ("Spread > 0   ∧   |z| ≤ 1",         "mid-cycle"),
            ("Spread > 0   ∧   z < −1",          "late-cycle"),
        ],
        interpretation="The most empirically validated recession "
            "predictor in macroeconomics. The 3-month yield tracks the "
            "current Fed Funds Rate; the 10-year blends future-rate "
            "expectations with the term premium. Inversion implies the "
            "market expects the Fed will need to cut rates materially — "
            "which only happens in recessions.",
        references="Estrella & Mishkin (1996, NBER) — the curve "
            "outperforms all single recession indicators 4–6 quarters "
            "ahead. NY Fed publishes an ongoing recession-probability "
            "model based solely on this spread.",
    )

    # ── 4. US_V1 ────────────────────────────────────────────────────────
    add_indicator_card(doc,
        ind_id="US_V1",
        full_name="VIX Term Structure (3-month minus 1-month)",
        eq_builder=EQ.eq_us_v1,
        data="CBOE 3-Month VIX (^VIX3M) and CBOE VIX (^VIX), yfinance.",
        lookback="Level-based regime with z-modulation in the positive "
                 "regime; z over 156 weeks.",
        regime_rows=[
            ("VIX3M − VIX < 0   (inversion)",   "stress"),
            ("Spread > 0   ∧   z < −1",          "complacency"),
            ("Spread > 0   ∧   |z| ≤ 1",         "normal"),
        ],
        interpretation="Implied volatility term structures are normally "
            "in contango — three-month vol exceeds one-month vol because "
            "uncertainty grows with horizon. Inversion is a hallmark of "
            "acute stress; investors are paying premium prices for "
            "immediate hedges. Inversion has preceded most major market "
            "dislocations of the last two decades.",
        references="Whaley (2009, JFM) on VIX term structure as a "
            "real-time fear gauge; Carr & Wu (2006) on the variance "
            "risk premium in inversions.",
    )

    # ── 5. FX_CMD1 ──────────────────────────────────────────────────────
    add_indicator_card(doc,
        ind_id="FX_CMD1",
        full_name="Copper / Gold (Global Growth Barometer)",
        eq_builder=EQ.eq_fx_cmd1,
        data="Copper futures (HG=F) and gold futures (GC=F), yfinance.",
        lookback="156-week rolling z-score of the log-ratio.",
        regime_rows=[
            ("z > +1",        "pro-growth — OW cyclicals, EM commodities"),
            ("|z| ≤ 1",       "neutral"),
            ("z < −1",        "growth-scare — OW gold, defensive duration"),
        ],
        interpretation="Copper is the industrial bellwether (about 60% "
            "of demand from construction, electrical, machinery); gold "
            "is the safe-haven. The ratio answers a single question: is "
            "the marginal capital seeking productive investment or "
            "protecting itself? The ratio leads the 10-year Treasury "
            "yield by six to twelve months.",
        references="Erb & Harvey (2006, FAJ) on copper-PMI linkage; "
            "Gundlach (DoubleLine commentary) popularising the ratio "
            "as a 10-year-yield leading indicator.",
    )

    # ── 6. GL_PMI1 ──────────────────────────────────────────────────────
    add_indicator_card(doc,
        ind_id="GL_PMI1",
        full_name="Equal-weight Global Manufacturing Survey Composite",
        eq_builder=EQ.eq_gl_pmi1,
        data="ISM Manufacturing PMI (DB.nomics), Eurostat Industrial "
             "Confidence (DB.nomics), UK Business Confidence (FRED OECD "
             "mirror), German ifo Industry composite (ifo direct), and "
             "Japan Tankan Large Manufacturer DI (BoJ Time-Series).",
        lookback="Each component independently z-scored against its own "
                 "156-week window; the composite is the equal-weight "
                 "average of available z-scores.",
        regime_rows=[
            ("z > +1",        "global-expansion — broad survey-led growth"),
            ("|z| ≤ 1",       "neutral"),
            ("z < −1",        "global-contraction — broad survey-led decline"),
        ],
        interpretation="A free-data substitute for the proprietary "
            "JP Morgan Global Manufacturing PMI. The composite "
            "degrades gracefully when individual components are "
            "missing — only the available z-scores are averaged — "
            "so the indicator continues to compute through "
            "publisher-side outages of any single component.",
        references="Marquette (1992) on ISM new orders leading "
            "industrial production; methodological convention "
            "follows the OECD Composite Leading Indicator framework "
            "of equal-weight z-score aggregation.",
    )


# ════════════════════════════════════════════════════════════════════════
#         SECTION 7 — REGIME, FORWARD-REGIME & TRAJECTORY CLASSIFICATION
# ════════════════════════════════════════════════════════════════════════

def add_section_7(doc):
    add_section_heading(doc, 7,
        "Regime, Forward-Regime & Z-Score Trajectory Classification")

    add_paragraph(doc,
        "Each indicator emits four classification labels alongside its "
        "raw and z-scored values: a current regime, a forward regime, "
        "a z-score trajectory diagnostic, and a cycle-timing tag. "
        "Together these labels carry the analytical interpretation of "
        "the indicator at every time step. This section documents how "
        "each label is computed."
    )

    add_subsection_heading(doc, "7.1", "Regime classification")

    add_paragraph(doc,
        "The default regime classifier is a three-bucket helper "
        "applied to the z-score:"
    )

    add_eq_display(doc, EQ.regime_3bucket, eq_number="7.1")

    add_paragraph(doc,
        "Each indicator supplies its own labels for L₊, L₀, and L₋. "
        "For US_G1 these are pro-growth / neutral / defensive; for "
        "FX_CMD1 they are pro-growth / neutral / growth-scare; for "
        "US_M2 they are abundant-liquidity / neutral / tight-liquidity. "
        "The labels are not comparable across indicators — they are "
        "interpretation aids rather than universal categories — but "
        "the underlying ±1 z-threshold convention is uniform."
    )

    add_subsection_heading(doc, "7.2", "Indicator-specific overrides")

    add_paragraph(doc,
        "A minority of indicators apply a level override to the "
        "default classifier. The overrides exist for two reasons: "
        "structural levels carry signal that a rolling z-score would "
        "lose by adapting to regime drift, and certain economic "
        "conditions (yield-curve inversion, vol-curve inversion, "
        "credit spread above 800 basis points) are absolute states "
        "that should not be relativised."
    )

    add_paragraph(doc,
        "Three patterns recur: yield-curve inversion (US_R1, US_R2, "
        "US_R3) — any negative spread overrides the z-score to "
        "recession-watch; volatility-curve inversion (US_V1) — any "
        "negative VIX3M − VIX overrides to stress; and the five-regime "
        "credit-spread classifier (US_Cr2) which couples raw basis-"
        "point levels with z-score thresholds across five distinct "
        "regimes."
    )

    add_subsection_heading(doc, "7.3", "Forward regime")

    add_paragraph(doc,
        "The forward regime captures the indicator's recent trajectory "
        "rather than its current level. It is computed from the slope "
        "of the z-score over the trailing eight weeks:"
    )

    add_eq_display(doc, EQ.fwd_slope, eq_number="7.2")

    add_paragraph(doc,
        "and classified using ±0.15 thresholds:"
    )

    add_eq_display(doc, EQ.fwd_regime, eq_number="7.3")

    add_paragraph(doc,
        "The eight-week window is short enough to capture genuine "
        "trajectory shifts and long enough to dampen weekly noise. "
        "The 0.15 threshold corresponds to a z-score change of about "
        "1.2 over the eight weeks — small enough to register typical "
        "regime transitions and large enough to filter out random walks. "
        "Indicators flagged as naturally leading in the registry receive "
        "a [leading] suffix on their forward-regime label, marking that "
        "their trajectory is itself a forward-looking signal rather "
        "than a coincident one."
    )

    add_subsection_heading(doc, "7.4", "Z-score trajectory diagnostic")

    add_paragraph(doc,
        "A separate diagnostic classifies the recent z-score trajectory "
        "against three lookbacks (1, 4, and 13 weeks) and the trailing "
        "13-week peak in absolute z-score. Four classes are recognised."
    )

    add_paragraph(doc,
        "Intensifying — z is rising in magnitude vs both 1- and 4-week "
        "lags, and is near the 13-week peak in absolute value:",
        size=10.5, space_after=2)
    add_eq_display(doc, EQ.trend_intensifying, eq_number="7.4")

    add_paragraph(doc,
        "Fading — current absolute z is meaningfully below its 4-week "
        "lag:",
        size=10.5, space_after=2)
    add_eq_display(doc, EQ.trend_fading, eq_number="7.5")

    add_paragraph(doc,
        "Reversing — z has changed sign vs the 4-week lag, from a "
        "previous reading already meaningful in magnitude:",
        size=10.5, space_after=2)
    add_eq_display(doc, EQ.trend_reversing, eq_number="7.6")

    add_paragraph(doc,
        "Stable — none of the above. The trajectory is purely diagnostic; "
        "it does not feed the regime call. Its purpose is to give a "
        "reader of the snapshot output a one-glance read on whether "
        "an indicator's signal is strengthening, weakening, or "
        "turning over without re-examining the raw z-score history. "
        "The four supplementary z-score samples (one, four, and "
        "thirteen weeks ago, plus the thirteen-week peak in absolute "
        "value) are also exposed in the snapshot CSV so a reader can "
        "verify the classification."
    )

    add_subsection_heading(doc, "7.5", "Cycle-timing classification")

    add_paragraph(doc,
        "Each indicator carries a cycle-timing tag of Leading, "
        "Coincident, or Lagging in the indicator-library CSV. The "
        "tag is determined editorially against academic and "
        "practitioner literature for that indicator's referent — "
        "yield-curve inversion is Leading by approximately twelve "
        "months, ISM new orders is Leading by approximately six "
        "weeks, NBER recession dates are Lagging. Of the ninety-two "
        "indicators currently in the library, ninety are Leading, "
        "two are Coincident, and none are Lagging — a deliberate "
        "skew, since the project's analytical purpose is "
        "forward-looking and the indicator selection has been biased "
        "accordingly."
    )

    add_callout(doc,
        "The four labels read together",
        "A high-conviction regime call combines all four signals: "
        "the current regime is on the right side of the bucket "
        "boundary, the forward regime carries [leading] and reads "
        "improving, the z-score trajectory is intensifying, and the "
        "cycle-timing tag is Leading. The opposite — current regime "
        "deep in one bucket while the forward regime contradicts and "
        "the trajectory is reversing — is the early-warning state "
        "that a regime transition is in progress. Regime-based asset "
        "allocation work depends on this multi-signal reading."
    )


# ════════════════════════════════════════════════════════════════════════
#                  SECTION 8 — OPERATIONAL SCAFFOLDING
# ════════════════════════════════════════════════════════════════════════

def add_section_8(doc):
    add_section_heading(doc, 8, "Operational Scaffolding")

    add_paragraph(doc,
        "Daily operation of a multi-source data pipeline produces a "
        "non-trivial volume of failure modes: HTTP errors, dead "
        "tickers, schema drift between code and registry, silent "
        "publisher freezes that the Friday-spine forward-fill would "
        "otherwise mask, and retroactive source-side history "
        "truncation. An integrated audit subsystem detects all of "
        "these in a single daily report and surfaces them through a "
        "perpetual GitHub Issue thread. A writeback step automates "
        "the most common remediation (retiring persistently dead "
        "tickers). A separate operator-gated utility synchronises "
        "history files with their source-of-truth registries when "
        "rows are deliberately removed."
    )

    add_subsection_heading(doc, "8.1", "The daily integrated audit")

    add_paragraph(doc,
        "The audit runs as the post-fetch, post-explorer step of the "
        "GitHub Actions workflow. It produces two artefacts: a full "
        "plaintext report (data_audit.txt) committed alongside the "
        "data CSVs, and a Markdown body (audit_comment.md) posted as "
        "a comment to the perpetual daily-audit GitHub Issue. The "
        "first line of the comment is always one of:"
    )

    add_bullet(doc, "## Daily audit — YYYY-MM-DD — ALL CLEAN")
    add_bullet(doc,
        "## Daily audit — YYYY-MM-DD — N ISSUES "
        "(X fetch errors, Y static-check failures, Z stale series, …)")

    add_paragraph(doc,
        "Detail follows in collapsible Markdown sections; row counts "
        "per category are capped for readability with full detail in "
        "the plaintext report. The audit is a strict warning channel — "
        "it always exits cleanly, never failing the workflow."
    )

    add_subsection_heading(doc, "8.2", "Section A — fetch outcomes")

    add_paragraph(doc,
        "Section A scrapes the freshly-captured pipeline.log for known "
        "failure patterns: HTTP 4xx/5xx errors that survived the "
        "exponential backoff, yfinance ‘possibly delisted’ warnings, "
        "and ‘Quote not found for symbol’ messages. yfinance suspects "
        "are cross-checked against the latest non-empty row of the "
        "comp-pipeline history file to filter transient warnings — a "
        "ticker that returned no data for one fetch but shows up in "
        "the latest history row was clearly a transient failure, not "
        "a delisting. Retried-then-recovered transients are filtered "
        "out by matching only the ‘— skipping’ suffix in the log line."
    )

    add_subsection_heading(doc, "8.3", "Section B — static checks")

    add_paragraph(doc,
        "Section B runs a battery of consistency checks against the "
        "registry CSVs and the code that consumes them. Each is a "
        "local operation requiring no network access:"
    )

    add_bullet(doc,
        "Country-code orphans: every code referenced in the per-source "
        "libraries must exist in macro_library_countries.csv.")
    add_bullet(doc,
        "Indicator-id uniqueness: no two rows in "
        "macro_indicator_library.csv share an id.")
    add_bullet(doc,
        "Calculator registration: every id in the indicator library "
        "must be registered as a calculator in compute_macro_market.py.")
    add_bullet(doc,
        "Column-existence: every literal passed to _get_col(...) by a "
        "calculator must resolve to a column in the unified macro hist.")
    add_bullet(doc,
        "Registry drift across all three history-vs-library pairs "
        "(comp / macro_economic / macro_market): every column in each "
        "history file must trace back to a row in its source-of-truth "
        "library. Orphan columns are reported with the recommended "
        "remediation: ‘run python library_sync.py --confirm’.")

    add_subsection_heading(doc, "8.4", "Section C — value-change staleness")

    add_paragraph(doc,
        "Section C addresses the failure mode that the Friday-spine "
        "forward-fill would otherwise mask: a publisher silently "
        "stopping its updates while the most-recent observation continues "
        "to populate every Friday slot via forward-fill. The audit "
        "walks every column of the unified macro hist, finds the last "
        "date on which the value actually changed (rather than the "
        "last non-null cell), and computes the age in days:"
    )

    add_eq_display(doc, EQ.staleness_age, eq_number="8.1")

    add_paragraph(doc,
        "This age is then bucketed against a per-frequency tolerance T "
        "loaded from data/freshness_thresholds.csv (5 days for daily "
        "series, 10 for weekly, 45 for monthly, 120 for quarterly, "
        "540 for annual). Per-row overrides on the per-source library "
        "files allow individual series with longer-than-normal "
        "publisher lag to be calibrated upward. The classification is:"
    )

    add_eq_display(doc, EQ.staleness_classification, eq_number="8.2")

    add_paragraph(doc,
        "FRESH series are silent in the audit. STALE series receive "
        "a one-line note. EXPIRED series are surfaced in the headline "
        "issue count and listed individually in the report. A 2026-04-29 "
        "bulk pass widened tolerances on 48 rows where the 30-50 day "
        "publisher cadence on monthly series was producing benign STALE "
        "noise; the cluster-to-override mapping is recorded in "
        "technical_manual.md §9.8 as the durable record of what was "
        "chosen and why."
    )

    add_subsection_heading(doc, "8.5", "Section D — history preservation")

    add_paragraph(doc,
        "Section D, added 30 April 2026 in response to the ICE BofA "
        "truncation event, reports per-history-file row counts (live "
        "file, sister x-file, and union) and the date range of each. "
        "Two specific alerts fire: (a) any ICE-BofA-bearing live file "
        "with no sister x-file, and (b) sister rows being a strict "
        "subset of live rows, which would indicate a regression in "
        "the history-preservation writer. The mechanics of the "
        "preservation rule are documented in §5.9; Section D is the "
        "monitoring channel for it."
    )

    add_subsection_heading(doc, "8.6", "Library writeback")

    add_paragraph(doc,
        "audit_writeback.py runs immediately after the audit. It "
        "parses Section A's yfinance-dead list, maintains a "
        "per-ticker dead-streak counter in "
        "data/yfinance_failure_streaks.csv, and flips a row's "
        "validation_status from CONFIRMED to UNAVAILABLE once its "
        "streak reaches fourteen consecutive days. A single-line "
        "summary of any flips is appended to audit_comment.md so "
        "the daily GitHub-Issue notification surfaces them."
    )

    add_paragraph(doc,
        "The fourteen-day threshold is an operationally-tuned "
        "compromise. It is long enough that genuine transient outages "
        "(a weekend yfinance throttle, a one-day FRED maintenance "
        "window) do not trigger spurious retirement, and short enough "
        "that a permanently delisted ticker is removed from the "
        "fetch budget within a fortnight. Manual override always "
        "wins — re-setting validation_status to CONFIRMED after a "
        "real fix restarts the streak at zero on the next day's run."
    )

    add_subsection_heading(doc, "8.7", "Library-sync utility")

    add_paragraph(doc,
        "library_sync.py is the operator-gated companion to the "
        "registry-drift static check. When a row is deliberately "
        "removed from a registry CSV, the corresponding column in "
        "the relevant history file becomes orphaned. The utility "
        "covers three history-versus-library pairs: the comp pipeline "
        "(index_library.csv against market_data_comp_hist.csv), the "
        "macro-economic layer (the union of seven library CSVs against "
        "macro_economic_hist.csv), and the macro-market layer "
        "(macro_indicator_library.csv against macro_market_hist.csv "
        "with each indicator id producing four columns). For each "
        "orphan column, the existing data is archived to "
        "data/_archived_columns/ before the column is dropped from "
        "the live history file. Default mode is dry-run; the operator "
        "passes --confirm to apply."
    )

    add_callout(doc,
        "Why archive rather than delete?",
        "An orphan column is the result of a deliberate operator "
        "edit to the registry — a series being retired, a ticker "
        "being replaced with a proxy. Archiving preserves the "
        "historical data so that subsequent analytical work, or a "
        "reversal of the decision, can recover it without a "
        "round-trip to the original publisher. The archived files "
        "are tiny and version-controlled; the storage cost is "
        "negligible against the analytical cost of irreversible "
        "loss."
    )


# ════════════════════════════════════════════════════════════════════════
#         SECTION 9 — THE INDICATOR EXPLORER & 17-CONCEPT TAXONOMY
# ════════════════════════════════════════════════════════════════════════

def add_section_9(doc):
    add_section_heading(doc, 9, "The Indicator Explorer & Concept Taxonomy")

    add_paragraph(doc,
        "The CSV outputs are the canonical record of what the pipeline "
        "produced on any given day. They are not, however, the artefact "
        "from which a human analyst typically reads the data. That role "
        "is filled by the Indicator Explorer — a self-contained HTML "
        "page rebuilt nightly from the freshly-written CSVs and "
        "committed to git alongside them, additionally published to "
        "GitHub Pages so any browser can load it without authentication."
    )

    add_subsection_heading(doc, "9.1", "Layout")

    add_paragraph(doc,
        "The explorer is organised around a three-section sidebar. The "
        "first section, Macro Market Indicators, lists the ninety-two "
        "Phase E composites. The second, Economic Data, surfaces every "
        "raw macro series across the ten data sources in a single "
        "merged tree. The third, Market Data, exposes the comp-pipeline "
        "price series with a Local-currency / USD variant toggle."
    )

    add_paragraph(doc,
        "Each of the first two sections supports two grouping modes: "
        "By Region (the indicator's regional bucket — US, UK, Europe, "
        "Japan, Asia, Global, FX & Commodities) and By Concept (against "
        "the seventeen-concept taxonomy described below). Mode "
        "switching preserves the user's checked-state for plotted "
        "series, so the analyst can move between an asset-class "
        "lens and a thematic lens without losing position."
    )

    add_subsection_heading(doc, "9.2", "The 17-concept taxonomy")

    add_paragraph(doc,
        "The taxonomy was developed on 2026-04-28 to give the "
        "indicator library a thematic grouping orthogonal to its "
        "regional one. It applies uniformly across both the Phase E "
        "composites and the raw macro series — the same concept "
        "labels appear in both sidebar trees."
    )

    rows = [
        ("Equity",                "Cyclicals, defensives, factor / size / sector rotation"),
        ("Rates / Yields",        "Yield-curve slopes, real rates, policy paths"),
        ("Credit / Spreads",      "IG OAS, HY OAS, sovereign spreads, BTP-Bund"),
        ("Inflation",             "CPI, PPI, breakevens, inflation expectations"),
        ("Sentiment / Survey",    "PMIs, OECD CLI/BCI/CCI, ifo, Tankan, ZEW"),
        ("Leading Indicators",    "OECD CLI composites, Conference Board LEI"),
        ("Growth",                "GDP, industrial production, retail volumes"),
        ("Labour",                "Unemployment, jobless claims, JOLTS, payrolls"),
        ("Consumer",              "Consumer confidence, retail sales, dispositions"),
        ("Housing",               "Building permits, mortgage spread, NAHB"),
        ("Manufacturing",         "ISM Manufacturing, IIP, capacity utilisation"),
        ("External / Trade",      "Trade balances, current account, FX reserves"),
        ("Money / Liquidity",     "M2, financial conditions, NFCI"),
        ("Cross-Asset",           "SPY/GOVT, ACWI/AGG, copper/gold cross signals"),
        ("FX",                    "Dollar index, EM currency baskets, single-pair"),
        ("Volatility",            "VIX, VIX term structure, MOVE, MOVE/VIX ratio"),
        ("Momentum",              "Trend-following, breadth, dual momentum"),
    ]
    add_shaded_table(doc,
        headers=["Concept", "Coverage"],
        rows=rows,
        col_widths=[Inches(1.6), Inches(4.6)],
    )

    add_subsection_heading(doc, "9.3", "Filtering and visualisation")

    add_paragraph(doc,
        "The explorer carries four orthogonal filters that compose "
        "freely: a free-text search box, the market-data variant "
        "toggle (Local / USD), an L/C/G chip filter on cycle-timing, "
        "and a country dropdown that surfaces the twelve canonical "
        "country codes. Empty groups collapse automatically when "
        "filters reduce them to zero rows. The country dropdown is "
        "populated from macro_library_countries.csv at build time, "
        "so additions to the country registry propagate to the "
        "explorer with no JavaScript edit."
    )

    add_paragraph(doc,
        "Each indicator panel renders three things together: the raw "
        "series as a line, the rolling z-score as a secondary axis, "
        "and a coloured regime strip running along the time axis. "
        "Four colours encode the regime palette: green for the "
        "positive regime, red for the negative, gold for amber / "
        "transitional regimes, and grey for neutral. The forward-"
        "regime label appears as a coloured badge alongside the "
        "current regime; an L, C, or G badge marks the indicator's "
        "cycle-timing tag."
    )

    add_paragraph(doc,
        "A custom PNG snapshot button composites the chart title, "
        "the Plotly chart image, the legend entries, and the regime "
        "colour key onto a single canvas — useful for pasting an "
        "indicator's current state into a research note without a "
        "screenshot tool."
    )


# ════════════════════════════════════════════════════════════════════════
#                  SECTION 10 — OUTPUTS & INTEGRATION
# ════════════════════════════════════════════════════════════════════════

def add_section_10(doc):
    add_section_heading(doc, 10, "Outputs & Integration Points")

    add_paragraph(doc,
        "The pipeline emits five categories of artefact each day. "
        "Three are data products consumed by humans or downstream "
        "code; two are operational artefacts consumed by the audit "
        "subsystem and the GitHub Issue notification channel."
    )

    rows = [
        ("Google Sheets workspace",
         "Seven active tabs in a single spreadsheet (ID 12nKIUGHz5...). "
         "Authoritative human-readable surface for the daily snapshot, "
         "history, and indicator outputs."),
        ("Version-controlled CSVs in data/",
         "One CSV per Sheets tab, plus the audit and ledger files. "
         "Programmatic consumption surface; every output is pinned to "
         "a Git commit with date-stamped message."),
        ("Indicator Explorer (HTML)",
         "Self-contained HTML + JavaScript bundle, committed to "
         "docs/indicator_explorer.html and published to GitHub Pages. "
         "Reads the freshly-written CSVs for every rebuild."),
        ("Pipeline log (pipeline.log)",
         "Captured stdout and stderr of the daily Python run, "
         "committed to the repo for diagnostic recovery without "
         "needing to download workflow artefacts."),
        ("Daily audit (data_audit.txt + audit_comment.md)",
         "Plaintext report committed to the repo, plus Markdown body "
         "posted as a comment to the perpetual daily-audit GitHub "
         "Issue. First line is the one-sentence ALL CLEAN / N ISSUES "
         "summary; GitHub native notifications email the operator on "
         "every comment."),
    ]
    add_shaded_table(doc,
        headers=["Artefact", "Purpose"],
        rows=rows,
        col_widths=[Inches(2.0), Inches(4.5)],
    )

    add_subsection_heading(doc, "10.1", "Downstream consumers")

    add_paragraph(doc,
        "One known downstream consumer reads the pipeline output today: "
        "a separate Windows-local script (trigger.py) running on a "
        "06:15 London cron, which reads the market_data tab via Google "
        "Sheets CSV export. Because trigger.py is the project's only "
        "production-critical reader of the simple-pipeline output, the "
        "market_data tab is included in the SHEETS_PROTECTED_TABS "
        "frozenset and is verified untouched by every writer in the "
        "pipeline."
    )

    add_paragraph(doc,
        "Beyond trigger.py, the data layer is designed to be a generic "
        "integration surface. Any downstream workflow — manual analysis "
        "in Sheets, ad-hoc Python notebooks against the committed CSVs, "
        "the planned regime-allocation back-test tooling, or research "
        "notes — consumes the same artefacts on the same schedule "
        "without coordination with the pipeline. The architectural "
        "rule is that downstream code reads but never writes; "
        "operational state lives only in the registry CSVs and the "
        "audit ledger files."
    )

    add_callout(doc,
        "Integration contract",
        "Every output file is fully regenerated each daily run; "
        "downstream code should read by date column and not by row "
        "index. Metadata prefix rows on history files (10 rows on the "
        "comp history, 14 on the macro history) must be skipped "
        "explicitly. The committed CSVs are the only stable contract "
        "— Sheets tab GIDs are stable but their column ordering is "
        "not strictly guaranteed across schema changes."
    )


# ════════════════════════════════════════════════════════════════════════
#               SECTION 11 — COVERAGE VS REFERENCE BASELINE
# ════════════════════════════════════════════════════════════════════════

def add_section_11(doc):
    add_section_heading(doc, 11, "Coverage Today vs Reference Baseline")

    add_paragraph(doc,
        "A standalone reference document — Macro Market Indicators "
        "Reference.docx — enumerates 206 macro and market indicators "
        "across six regions with explicit Leading / Coincident / "
        "Lagging classifications. That document is the demand-side "
        "specification against which the pipeline's coverage is "
        "evaluated. The current state, refreshed 1 May 2026 after "
        "Stage F community-ticker additions, is recorded below."
    )

    add_subsection_heading(doc, "11.1", "Per-region coverage")

    add_paragraph(doc,
        "Captured rows are those for which a free programmatic "
        "source is wired and producing a clean weekly observation. "
        "Partial captures use a proxy series (an ETF total-return "
        "instead of a yield, for example, or a sub-component instead "
        "of the headline). Missing rows are awaiting either a free "
        "source or an architectural decision."
    )

    rows = [
        ("US",        "37",  "28", "2",  "7",  "78%"),
        ("UK",        "36",  "3",  "2",  "31", "11%"),
        ("Eurozone",  "36",  "12", "8",  "16", "44%"),
        ("Japan",     "35",  "6",  "3",  "26", "21%"),
        ("China",     "36",  "8",  "4",  "24", "28%"),
        ("Global",    "26",  "9",  "3",  "14", "40%"),
        ("Total",     "206", "66", "22", "118","37%"),
    ]
    fills = [None, COLOR_PINK, None, None, None, None, COLOR_HEADER]
    add_shaded_table(doc,
        headers=["Region", "Total", "Full", "Partial", "Missing", "% Captured"],
        rows=rows,
        row_fills=fills,
        col_widths=[Inches(1.4), Inches(0.8), Inches(0.8), Inches(0.9),
                    Inches(0.9), Inches(1.2)],
    )

    add_subsection_heading(doc, "11.2", "Coverage by concept × region")

    add_paragraph(doc,
        "Cell entries below are ‘full / total’ counts. Concepts "
        "selected are those most material to the regime-classification "
        "axes (growth, inflation, rates, credit, sentiment, labour). "
        "Two regional standouts deserve explicit attention: the United "
        "Kingdom is the largest single-region coverage gap at eleven "
        "per cent, with growth at zero out of eight rows wired; and "
        "the Phase E inflation indicator family is sparse — only two "
        "of the ninety-two composites carry the Inflation concept "
        "tag — which constrains any future regime classifier that "
        "uses inflation as an axis."
    )

    rows = [
        ("Inflation",     "2/2", "1/1", "1/1", "1/3", "2/2", "0/0", "7/9"),
        ("Growth",        "6/6", "0/8", "6/6", "3/7", "3/9", "0/1", "18/37"),
        ("Rates / Yields","3/3", "2/2", "3/3", "1/2", "2/2", "—",   "11/12"),
        ("Credit / Spread","4/5","0/4", "0/4", "0/1", "0/4", "1/3", "5/21"),
        ("Sentiment",     "2/6", "1/9", "7/11","1/9", "1/7", "1/2", "13/44"),
        ("Labour",        "6/6", "0/6", "2/5", "1/3", "1/2", "—",   "10/22"),
    ]
    add_shaded_table(doc,
        headers=["Concept", "US", "UK", "EZ", "JP", "CN", "Global", "Total"],
        rows=rows,
        col_widths=[Inches(1.4), Inches(0.6), Inches(0.6), Inches(0.6),
                    Inches(0.6), Inches(0.6), Inches(0.6), Inches(0.7)],
    )

    add_subsection_heading(doc, "11.3", "Source-fallback realisation")

    add_paragraph(doc,
        "Of nine FRED series carrying internationally-mirrored macro "
        "data that have stopped updating (forcing-function rows "
        "documented in Section 4.5), seven have been resolved through "
        "fallback chains — four at the tier-one level via DB.nomics "
        "IMF/IFS and Eurostat, and three at the tier-two level via "
        "the BoE IADB, the ECB Data Portal, the BoJ Time-Series API, "
        "and e-Stat. The remaining two — Chinese M2 and Chinese "
        "industrial production — are accepted gaps because the People's "
        "Bank of China and the National Bureau of Statistics publish "
        "no free programmatic interface. The forcing-function FRED "
        "rows are deliberately retained in the registry so the daily "
        "audit continues to surface them as expired; the actual data "
        "now flows through the fallback sources."
    )


# ════════════════════════════════════════════════════════════════════════
#               SECTION 12 — KNOWN LIMITATIONS & ACCEPTED GAPS
# ════════════════════════════════════════════════════════════════════════

def add_section_12(doc):
    add_section_heading(doc, 12, "Known Limitations & Accepted Gaps")

    add_paragraph(doc,
        "The free-data constraint imposes hard limits. Where a "
        "limitation has been investigated and judged irreducible, "
        "it is recorded as an accepted gap rather than an open "
        "task. The rationale is durable: ‘this gap exists because "
        "the source is paywalled / scraping is fragile / no free "
        "API exists’ is a stable conclusion, and recording it once "
        "saves repeated investigation."
    )

    add_subsection_heading(doc, "12.1", "Data gaps")

    rows = [
        ("China 10-year govt yield",
         "Proprietary",
         "ETF proxy via CBON (VanEck China AMC China Bond) for "
         "regime use; direct yield series remains unsourced."),
        ("Euro IG corporate effective yield",
         "FRED series 400s; no free aggregator",
         "ETF proxy via IEAC.L (iShares Core EUR IG Corp Bond) for "
         "regime use; EU_Cr1 returns n/a until a corp-yield series "
         "is wired. EU_Cr2 (Euro HY) covered separately."),
        ("Caixin China Manufacturing PMI",
         "S&P Global proprietary",
         "Substitute via CN_PMI1 (OECD BCI for China)."),
        ("ZEW Economic Sentiment",
         "ZEW Mannheim licences the archive",
         "Substitute via DE_IFO1 + DEU_BUS_CONF."),
        ("au Jibun Bank Japan Manufacturing PMI",
         "S&P Global proprietary",
         "Functionally resolved via JP_TANKAN1 (BoJ Tankan Large "
         "Mfg DI) wired through sources/boj.py."),
        ("OECD CLI for EA19 / CHE",
         "Not published by OECD",
         "Eurozone CLI proxied as DEU+FRA equal-weight average."),
        ("Atlanta Fed GDPNow",
         "Web scrape only — no clean API",
         "Excluded; tracked in forward planning as a future "
         "Stage-C extension."),
        ("ICE BofA pre-2023 history on FRED",
         "ICE Data licence change April 2026 — rolling 3-year window",
         "Architectural mitigation via per-column floor-advancement "
         "detection and append-only sister files (Section 5.9)."),
    ]
    add_shaded_table(doc,
        headers=["Gap", "Reason", "Resolution / proxy"],
        rows=rows,
        col_widths=[Inches(1.6), Inches(1.7), Inches(3.2)],
    )

    add_subsection_heading(doc, "12.2", "Permanently unavailable instruments")

    add_paragraph(doc,
        "A small number of instruments are permanently unavailable "
        "via yfinance and have been removed from the comp-pipeline "
        "library. Russian instruments — IMOEX.ME and RTSI.ME — return "
        "data only through mid-2022 / mid-2024, owing to sanctions; "
        "their later history is not retrievable. WisdomTree's Chinese "
        "Yuan ETF (CYB) was delisted in December 2023 and has been "
        "replaced by CNYB.L. The US Dollar Index (DX-Y.NYB) returns "
        "data only from 2008 via yfinance and FRED's DTWEXBGS is used "
        "for longer histories where required by an indicator "
        "calculator."
    )

    add_subsection_heading(doc, "12.3", "Excluded sources")

    add_paragraph(doc,
        "Several candidate sources have been investigated and "
        "deliberately excluded. Investing.com and Trading Economics "
        "are protected by Cloudflare and explicitly forbid scraping "
        "in their terms of service. The Financial Modeling Prep API "
        "moved its economic-calendar endpoint behind a paywall in "
        "August 2025. The S&P Global PMI series and any series "
        "originating from a JP Morgan or Goldman Sachs proprietary "
        "index are licence-only. Bloomberg-terminal-only series — "
        "the canonical Goldman Sachs Financial Conditions Index, "
        "the Bloomberg Commodity Index sub-components — are excluded "
        "for the same reason. Where a free proxy exists (the Chicago "
        "Fed NFCI for the Goldman FCI; the equal-weight Global PMI "
        "composite for JP Morgan's), it is used. Where no proxy "
        "exists, the indicator is recorded as an accepted gap and not "
        "re-investigated."
    )

    add_subsection_heading(doc, "12.4", "Methodological caveats")

    add_paragraph(doc,
        "Two methodological caveats apply to every output of the "
        "system and merit explicit statement:"
    )

    add_bullet(doc,
        "Forward-fill into Friday slots is a deliberate choice that "
        "supports cross-series alignment but obscures intra-month "
        "publisher behaviour. The value-change staleness audit "
        "addresses the silent-freeze failure mode but cannot recover "
        "intra-month timing detail that a native-frequency store "
        "would carry. A multi-frequency rebuild is on the forward "
        "plan but is not in scope for this report.")
    add_bullet(doc,
        "Z-score thresholds at ±1 are convenient and broadly "
        "calibrated against the academic and practitioner literature "
        "for each indicator family, but they are not optimised "
        "against any specific objective function. A regime-aware "
        "threshold tuning step is on the forward plan; until it "
        "lands, the current thresholds are heuristic conventions "
        "rather than statistically derived cut-points.")


# ════════════════════════════════════════════════════════════════════════
#                       SECTION 13 — SUMMARY OF STATUS
# ════════════════════════════════════════════════════════════════════════

def add_section_13(doc):
    add_section_heading(doc, 13, "Summary of Status")

    add_paragraph(doc,
        "The system as it stands is in production. The daily run has "
        "executed cleanly on every market-day since the unified "
        "macro-economic coordinator landed on 23 April 2026. The "
        "ten-source data layer is comprehensive against the free-"
        "source ceiling, with seven of the nine identified "
        "forcing-function gaps closed through tier-one and tier-two "
        "fallback chains and the remainder accepted as residual "
        "limits of free programmatic data. The composite-indicator "
        "layer of ninety-two indicators is feature-complete — z-score, "
        "regime, forward-regime, trajectory, and cycle-timing labels "
        "all surface in the daily snapshot — and the operational "
        "scaffolding around it is mature: a four-section integrated "
        "audit, an automated writeback for dead tickers, an "
        "operator-gated library-sync utility, and a perpetual "
        "GitHub Issue thread for the daily heartbeat."
    )

    add_paragraph(doc,
        "The substrate is documented across three working "
        "documents: technical_manual.md (the authoritative record "
        "of the current code state, ~1,360 lines), forward_plan.md "
        "(the architectural rules, the project chronology, and the "
        "outstanding work queue, ~810 lines), and indicator_manual.md "
        "(the indicator-by-indicator reference with economic "
        "rationale and academic citations, ~1,850 lines). This "
        "report is a self-contained synthesis of the substrate and "
        "the design decisions behind it."
    )

    add_paragraph(doc,
        "The scope of this report is strictly retrospective. The "
        "regime-based asset allocation work that the present "
        "infrastructure is intended to support is out of scope by "
        "design — it is the subject of subsequent work, not of "
        "this record."
    )


# ════════════════════════════════════════════════════════════════════════
#                APPENDIX A — MATHEMATICAL NOTATION REFERENCE
# ════════════════════════════════════════════════════════════════════════

def add_appendix_a(doc):
    doc.add_page_break()
    h = doc.add_heading("Appendix A.  Mathematical Notation Reference", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(8)

    add_paragraph(doc,
        "This appendix consolidates the symbols and conventions used "
        "in the equations of Sections 4 through 8. Equation labels "
        "are reproduced as cross-references to their first appearance "
        "in the body."
    )

    add_subsection_heading(doc, "A.1", "Symbol conventions")

    rows = [
        ("t",                       "Time index, weekly (Friday close) unless noted otherwise"),
        ("τ",                       "Reference time strictly less than or equal to t"),
        ("x_t, y_t, a_t, b_t",      "Generic series values at time t"),
        ("r_local, r_FX, r_USD",    "Local-currency, FX, and USD-adjusted returns"),
        ("z_t",                     "Z-score at time t (3-year rolling)"),
        ("μ_W(t), σ_W(t)",          "Window mean and standard deviation over the trailing W observations"),
        ("W",                       "Z-score window length (W = 156 weeks)"),
        ("β",                       "Eight-week z-score slope (forward-regime regressor)"),
        ("I_t",                     "Composite indicator value at time t"),
        ("N_t, D_t",                "Numerator and denominator series for log-ratio indicators"),
        ("S^F_t",                   "Friday-spine value of series S at time t"),
        ("L_+, L_0, L_-",           "Indicator-specific positive / neutral / negative regime labels"),
        ("T",                       "Per-frequency staleness tolerance (days)"),
        ("age(s)",                  "Days since the last value-change of series s"),
    ]
    add_shaded_table(doc,
        headers=["Symbol", "Meaning"],
        rows=rows,
        col_widths=[Inches(1.6), Inches(4.5)],
    )

    add_subsection_heading(doc, "A.2", "Typographic conventions")

    add_paragraph(doc,
        "Single letter symbols (r, x, t, z, β) are typeset italic. "
        "Multi-letter labels denoting categorical content (USD, local, "
        "FX, ticker codes such as XLY and XLP, instrument labels, "
        "regime labels) are typeset upright. Function names (log, sgn, "
        "max) are typeset upright. Operators ( =, +, −, <, >, ≤, ≥, "
        "∧, ∨, ≠) use Word's default mathematical typography. Angle "
        "brackets and the floor / ceiling notation are not used "
        "anywhere in the report."
    )

    add_subsection_heading(doc, "A.3", "Equation index")

    rows = [
        ("4.1", "USD-adjusted return identity",        "Section 4.4"),
        ("4.2", "LSE pence correction",                 "Section 4.4"),
        ("5.1", "Friday-spine alignment",               "Section 5.3"),
        ("6.1", "Rolling z-score",                      "Section 6.2"),
        ("6.2", "Window mean μ_W with W = 156",         "Section 6.2"),
        ("6.3", "Log-ratio template",                   "Section 6.3"),
        ("6.4", "Composite log-ratio template",         "Section 6.3"),
        ("6.5", "Arithmetic-difference template",       "Section 6.3"),
        ("6.6", "Year-on-year monthly transform",       "Section 6.3"),
        ("7.1", "Three-bucket regime classifier",       "Section 7.1"),
        ("7.2", "Forward-regime slope β",                "Section 7.3"),
        ("7.3", "Forward-regime classification",        "Section 7.3"),
        ("7.4", "Trajectory: intensifying",             "Section 7.4"),
        ("7.5", "Trajectory: fading",                   "Section 7.4"),
        ("7.6", "Trajectory: reversing",                "Section 7.4"),
        ("8.1", "Staleness age",                        "Section 8.4"),
        ("8.2", "Staleness band classification",        "Section 8.4"),
    ]
    add_shaded_table(doc,
        headers=["Eq.", "Description", "First reference"],
        rows=rows,
        col_widths=[Inches(0.6), Inches(3.7), Inches(1.7)],
    )


# ════════════════════════════════════════════════════════════════════════
#                          APPENDIX B — REFERENCES
# ════════════════════════════════════════════════════════════════════════

def add_appendix_b(doc):
    doc.add_page_break()
    h = doc.add_heading("Appendix B.  References", level=1)
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after  = Pt(8)

    add_paragraph(doc,
        "Academic and practitioner references cited in the body of "
        "the report. Where multiple citations exist for a single "
        "concept, the most authoritative or earliest is selected."
    )

    refs = [
        ("Altman (1968)",
         "Altman, E. I. (1968). Financial ratios, discriminant "
         "analysis and the prediction of corporate bankruptcy. "
         "Journal of Finance, 23(4), 589–609."),
        ("Borio & Lowe (2002)",
         "Borio, C. & Lowe, P. (2002). Asset prices, financial and "
         "monetary stability: exploring the nexus. BIS Working Paper "
         "No. 114."),
        ("Carr & Wu (2006)",
         "Carr, P. & Wu, L. (2006). A tale of two indices. Journal "
         "of Derivatives, 13(3), 13–29."),
        ("Duffie & Singleton (1999)",
         "Duffie, D. & Singleton, K. J. (1999). Modeling term "
         "structures of defaultable bonds. Review of Financial "
         "Studies, 12(4), 687–720."),
        ("Erb & Harvey (2006)",
         "Erb, C. B. & Harvey, C. R. (2006). The strategic and "
         "tactical value of commodity futures. Financial Analysts "
         "Journal, 62(2), 69–97."),
        ("Estrella & Mishkin (1996)",
         "Estrella, A. & Mishkin, F. S. (1996). The yield curve as "
         "a predictor of US recessions. NBER Working Paper No. 5379. "
         "Current Issues in Economics and Finance, 2(7)."),
        ("Faber (2007)",
         "Faber, M. T. (2007). A quantitative approach to tactical "
         "asset allocation. Journal of Wealth Management, 9(4), 69–79."),
        ("Fama & French (1989)",
         "Fama, E. F. & French, K. R. (1989). Business conditions "
         "and expected returns on stocks and bonds. Journal of "
         "Financial Economics, 25(1), 23–49."),
        ("Friedman (1957)",
         "Friedman, M. (1957). A Theory of the Consumption Function. "
         "Princeton University Press."),
        ("Marquette (1992)",
         "Marquette, J. F. (1992). Forecasting with the NAPM "
         "Manufacturing Survey. Business Economics, 27(4)."),
        ("Merton (1974)",
         "Merton, R. C. (1974). On the pricing of corporate debt: "
         "the risk structure of interest rates. Journal of Finance, "
         "29(2), 449–470."),
        ("Whaley (2009)",
         "Whaley, R. E. (2009). Understanding the VIX. Journal of "
         "Portfolio Management, 35(3), 98–105."),
        ("OECD CLI methodology",
         "OECD (2012). OECD System of Composite Leading Indicators. "
         "OECD Publishing — methodological reference for "
         "equal-weight z-score aggregation."),
        ("NY Fed Recession Probability Model",
         "Federal Reserve Bank of New York, Economic Indicators "
         "Calendar — ongoing publication of the recession-probability "
         "model based on the 10Y–3M term spread."),
    ]

    for short, full in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        r1 = p.add_run(f"{short}.  ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(full)
        r2.font.size = Pt(10)


# ── orchestration ───────────────────────────────────────────────────────

def main():
    doc = Document()
    set_document_properties(doc)
    setup_document(doc)
    configure_header_footer(doc)

    # Title page (first page — no header / footer applied due to
    # different-first-page setting in configure_header_footer).
    add_title_page(doc)

    # Table of contents (Word populates on first open).
    add_table_of_contents(doc)

    # Report body
    add_section_1(doc)
    add_section_2(doc)
    add_section_3(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)
    add_section_11(doc)
    add_section_12(doc)
    add_section_13(doc)

    # Appendices
    add_appendix_a(doc)
    add_appendix_b(doc)

    # Interim probe removed in Step 6 — every equation now lives in
    # its final position within the body of the report.

    out_dir  = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, OUT_FILENAME)
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
