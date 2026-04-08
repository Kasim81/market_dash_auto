"""Convert indicator_manual.md to indicator_manual.docx"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = "indicator_manual.md"
DOCX_PATH = "indicator_manual.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_inline_formatted(run_parent, text):
    """Add a paragraph with **bold** and `code` inline formatting."""
    # Split on bold (**...**) and inline code (`...`)
    pattern = re.compile(r'(\*\*.*?\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = run_parent.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = run_parent.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9)
        else:
            run_parent.add_run(part)

def strip_inline(text):
    """Strip markdown inline markers for plain-text contexts."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

# ── main parser ───────────────────────────────────────────────────────────────

def build_docx(md_path, docx_path):
    doc = Document()

    # ── page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── default body font ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip('\n')
        stripped = raw.strip()

        # blank line
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run()
            run.add_break()
            i += 1
            continue

        # headings
        if stripped.startswith('#'):
            m = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if m:
                level = len(m.group(1))
                text  = strip_inline(m.group(2))
                h = doc.add_heading(text, level=level)
                h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
                h.paragraph_format.space_after  = Pt(4)
                i += 1
                continue

        # block-quote (> ...)
        if stripped.startswith('>'):
            text = re.sub(r'^>\s?', '', stripped)
            text = strip_inline(text)
            p = doc.add_paragraph(style='Quote')
            p.text = text
            p.paragraph_format.left_indent = Inches(0.4)
            i += 1
            continue

        # table (lines starting with |)
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            # filter separator rows
            rows = [r for r in table_lines if not re.match(r'^\|[-| :]+\|$', r)]
            if not rows:
                continue
            # parse cells
            parsed = []
            for r in rows:
                cells = [c.strip() for c in r.strip('|').split('|')]
                parsed.append(cells)
            ncols = max(len(r) for r in parsed)
            t = doc.add_table(rows=len(parsed), cols=ncols)
            t.style = 'Table Grid'
            for ri, row in enumerate(parsed):
                for ci, cell_text in enumerate(row):
                    cell = t.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    add_inline_formatted(p, strip_inline(cell_text))
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                    # header row shading
                    if ri == 0:
                        set_cell_bg(cell, "D9E2F3")
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph()  # spacing after table
            continue

        # bullet list (- or *)
        if re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            add_inline_formatted(p, text)
            for run in p.runs:
                run.font.size = Pt(10.5)
            i += 1
            continue

        # italic-only paragraph (*...*) — used for section intros
        if stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**'):
            text = stripped[1:-1]
            p = doc.add_paragraph()
            r = p.add_run(strip_inline(text))
            r.italic = True
            r.font.size = Pt(10.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(6)
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(6)
        add_inline_formatted(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    import os
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    build_docx(MD_PATH, DOCX_PATH)
