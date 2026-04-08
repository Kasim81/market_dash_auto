"""Convert macro_market_cheat_sheet.md to a print-friendly .docx."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

MD = Path(__file__).parent / "macro_market_cheat_sheet.md"
OUT = Path(__file__).parent / "macro_market_cheat_sheet.docx"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def style_header_row(row):
    """Dark header row with white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, "2F5496")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(8)


def add_table(doc, header, rows):
    """Add a formatted table to the document."""
    ncols = len(header)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header
    for i, h in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.text = h.strip()
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    style_header_row(tbl.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data[:ncols]):
            cell = tbl.rows[r_idx + 1].cells[c_idx]
            cell.text = val.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
        # Alternating shading
        if r_idx % 2 == 1:
            for cell in tbl.rows[r_idx + 1].cells:
                set_cell_shading(cell, "D6E4F0")

    return tbl


def parse_md_table(lines, start):
    """Parse a markdown table starting at `start`. Returns (header, rows, end_idx)."""
    header_line = lines[start]
    header = [c.strip() for c in header_line.strip().strip("|").split("|")]
    # Skip separator line
    data_start = start + 2
    rows = []
    idx = data_start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cols = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        rows.append(cols)
        idx += 1
    return header, rows, idx


def add_formatted_paragraph(doc, text, style="Normal", bold=False):
    """Add a paragraph with inline bold/italic markdown formatting."""
    p = doc.add_paragraph(style=style)
    # Split on bold (**...**) and italic (*...*)
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
        else:
            run = p.add_run(part)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def main():
    lines = MD.read_text().splitlines()
    doc = Document()

    # Page setup — narrower margins for print
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(1)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines and horizontal rules
        if not stripped or stripped == "---":
            i += 1
            continue

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|---"):
            header, rows, end = parse_md_table(lines, i)
            add_table(doc, header, rows)
            doc.add_paragraph("")  # spacing
            i = end
            continue

        # Bullet list
        if stripped.startswith("- "):
            text = stripped[2:]
            add_formatted_paragraph(doc, text, style="List Bullet")
            i += 1
            continue

        # Numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            text = m.group(2)
            add_formatted_paragraph(doc, text, style="List Number")
            i += 1
            continue

        # Regular paragraph with formatting
        add_formatted_paragraph(doc, stripped)
        i += 1

    doc.save(str(OUT))
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
